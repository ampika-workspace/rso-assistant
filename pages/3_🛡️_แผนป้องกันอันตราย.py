"""
🛡️ Page 3 — แผนป้องกันอันตรายจากรังสี
อ้างอิง:
  - เอกสาร ปส.: เจ้าหน้าที่ความปลอดภัยทางรังสีกับการประเมินความปลอดภัยทางรังสี
  - IAEA Safety Reports Series No.47
  - NCRP Report No.151
  - กฎกระทรวงความปลอดภัยทางรังสี พ.ศ. 2561
"""
import streamlit as st
import math, sys, os, requests
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from utils.data import D_VALUES_TBQ, ISOTOPE_INFO, UNIT_TO_TBQ, classify_material, USE_TYPE_CATEGORY

st.set_page_config(page_title="แผนป้องกันอันตรายจากรังสี", page_icon="🛡️", layout="wide")
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@400;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Sarabun', sans-serif !important; }
</style>
""", unsafe_allow_html=True)

# ── ค่าคงที่และตารางอ้างอิง ────────────────────────────────────────────────

# HVL (cm) อ้างอิง NCRP 151 / IAEA SRS No.47
HVL_DATA = {
    ("low",       "ตะกั่ว (Pb)"):        0.017,
    ("low",       "คอนกรีต (2.35 g/cm³)"): 0.9,
    ("low",       "เหล็ก (Fe)"):          0.05,
    ("low",       "น้ำ/เนื้อเยื่อ"):      3.0,
    ("medium",    "ตะกั่ว (Pb)"):         0.6,
    ("medium",    "คอนกรีต (2.35 g/cm³)"): 6.1,
    ("medium",    "เหล็ก (Fe)"):          1.5,
    ("medium",    "น้ำ/เนื้อเยื่อ"):      10.0,
    ("high",      "ตะกั่ว (Pb)"):         1.2,
    ("high",      "คอนกรีต (2.35 g/cm³)"): 10.0,
    ("high",      "เหล็ก (Fe)"):          2.2,
    ("high",      "น้ำ/เนื้อเยื่อ"):      14.0,
    ("very_high", "ตะกั่ว (Pb)"):         1.4,
    ("very_high", "คอนกรีต (2.35 g/cm³)"): 12.0,
    ("very_high", "เหล็ก (Fe)"):          2.5,
    ("very_high", "น้ำ/เนื้อเยื่อ"):      16.0,
}

# TVL (cm) ≈ HVL × 3.32  (TVL = HVL × log₁₀(2) × 10)
def hvl_to_tvl(hvl): return hvl * math.log10(2) * 10

MATERIALS = ["ตะกั่ว (Pb)", "คอนกรีต (2.35 g/cm³)", "เหล็ก (Fe)", "น้ำ/เนื้อเยื่อ"]

ENERGY_GROUP_LABEL = {
    "low":       "Low (< 0.1 MeV) — Am-241, I-125, Pd-103",
    "medium":    "Medium (0.1–0.5 MeV) — Ir-192, Cs-137, Se-75",
    "high":      "High (0.5–1.5 MeV) — Co-60, Ra-226, Na-22",
    "very_high": "Very High (> 1.5 MeV) — Co-60 (1.25 MeV), La-140",
}

ISOTOPE_ENERGY_GROUP = {
    "Am-241": "low", "Am-241/Be": "low", "I-125": "low", "Pd-103": "low",
    "Tm-170": "low", "Co-57": "low", "Tc-99m": "low", "Gd-153": "low",
    "Tl-204": "low", "Pm-147": "low", "Cd-109": "low", "Cm-244": "low",
    "Ir-192": "medium", "Cs-137": "medium", "Se-75": "medium", "I-131": "medium",
    "Yb-169": "medium", "Mo-99": "medium", "Au-198": "medium", "In-111": "medium",
    "Ge-68": "medium", "Ru-106(Rh-106)": "high", "Ra-226": "high",
    "Na-22": "high", "Cs-134": "high", "Sr-90(Y-90)": "medium",
    "Co-60": "very_high", "La-140": "very_high",
    "H-3": "low", "P-32": "low", "Fe-55": "low", "Ni-63": "low",
    "Kr-85": "medium", "Pu-238": "low", "Po-210": "low", "Pu-239d/Be": "low",
    "Cf-252": "medium",
}

# Dose limits (ปส. / ICRP 103)
DOSE_LIMITS_WEEK = {
    "ผู้ปฏิบัติงานทางรังสี":  400,   # µSv/week → 20 mSv/yr
    "พื้นที่ควบคุม (Controlled Area)": 115,  # ~6 mSv/yr
    "พื้นที่ตรวจตรา (Supervised Area)": 20,  # 1 mSv/yr
    "ประชาชนทั่วไป": 20,             # 1 mSv/yr
}

# RAKR ของ isotope ต่างๆ (µGy·m²/MBq·h) — อ้างอิง IAEA SRS No.47
RAKR = {
    "Ir-192": 0.111,
    "Co-60":  0.306,
    "Cs-137": 0.077,
    "Ir-192 (HDR Brachytherapy)": 0.111,
}

# Occupancy factor T — อ้างอิง IAEA SRS No.47
OCCUPANCY_FACTORS = {
    "เต็มเวลา (ห้องทำงาน, ห้องรับรอง) T=1":         1.0,
    "ห้องน้ำ, บันได, ลิฟต์ (มีคนประจำ) T=1/4":    0.25,
    "ทางเดิน, ห้องน้ำสาธารณะ, บันได T=1/20":       0.05,
    "ที่จอดรถ, บริเวณกลางแจ้ง T=1/40":             0.025,
    "หน้าประตูห้องฉาย T=1/8":                       0.125,
    "ระเบียง, ทางเชื่อม T=1/5":                     0.2,
}

# Use factor U
USE_FACTORS = {
    "U=1 (ทุกทิศทาง / Brachytherapy / secondary beam)": 1.0,
    "U=1/4 (ผนังที่ถูกลำรังสีหลักเพียง 1/4 ของเวลา)":  0.25,
    "U=1/2 (ผนังด้านข้าง)":                             0.5,
    "U=0 (ผนังที่ไม่โดนลำรังสีหลัก)":                  0.0,
}

# ── Sidebar ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🏢 ข้อมูลองค์กร")
    org_name   = st.text_input("ชื่อองค์กร", placeholder="บริษัท / โรงพยาบาล ABC")
    license_no = st.text_input("เลขที่ใบอนุญาต", placeholder="4I0001/63F")
    rso_name   = st.text_input("ชื่อ RSO", placeholder="นายปรมาณู เพื่อสันติ")
    rso_phone  = st.text_input("โทรศัพท์ RSO", placeholder="08x-xxx-xxxx")
    plan_date  = st.date_input("วันที่จัดทำแผน")
    st.markdown("---")
    anthropic_key = st.text_input("🔑 Anthropic API Key", type="password")

st.markdown("# 🛡️ แผนป้องกันอันตรายจากรังสี")
st.caption(
    "อ้างอิง: เอกสาร ปส. (เจ้าหน้าที่ความปลอดภัยทางรังสีกับการประเมินความปลอดภัยทางรังสี) "
    "| IAEA SRS No.47 | NCRP No.151 | กฎกระทรวงความปลอดภัยทางรังสี พ.ศ. 2561"
)

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "☢️ แหล่งรังสี",
    "📐 กำบัง & พื้นที่",
    "🏥 ห้องฉายรังสี (NCRP/IAEA)",
    "👷 Dose ผู้ปฏิบัติงาน",
    "🤖 AI ร่างแผน & Export",
])

isotope_list = list(D_VALUES_TBQ.keys())
use_list = ["— ไม่พบในตาราง / คำนวณ A/D —"] + list(USE_TYPE_CATEGORY.keys())

# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    st.subheader("ข้อมูลแหล่งกัมมันตรังสี")
    c1, c2 = st.columns(2)
    with c1:
        isotope  = st.selectbox("ชนิดนิวไคลด์", isotope_list,
                                 index=isotope_list.index("Ir-192"))
        ca, cb   = st.columns(2)
        activity = ca.number_input("ค่ากัมมันตภาพ", min_value=0.0,
                                    value=5.5, format="%.4f")
        unit     = cb.selectbox("หน่วย", list(UNIT_TO_TBQ.keys()),
                                 index=list(UNIT_TO_TBQ.keys()).index("Ci"))
        use_type = st.selectbox("การใช้ประโยชน์", use_list)
        activity_tbq = activity * UNIT_TO_TBQ.get(unit, 1)
        st.caption(f"= **{activity_tbq:.4e} TBq**")
        if isotope in ISOTOPE_INFO:
            info = ISOTOPE_INFO[isotope]
            st.info(f"**{isotope}** | ครึ่งชีวิต: {info['halfLife']} | "
                    f"A₁={info['a1']} TBq | A₂={info['a2']} TBq")

    with c2:
        ut = use_type if use_type != "— ไม่พบในตาราง / คำนวณ A/D —" else ""
        class_result = classify_material(isotope, activity_tbq, ut)
        from utils.data import CATEGORY_COLORS, CATEGORY_BG
        if class_result.get("category"):
            cat = class_result["category"]
            st.markdown(f"""
            <div style="background:{CATEGORY_BG[cat]};border:2px solid {CATEGORY_COLORS[cat]};
                        border-radius:10px;padding:14px;">
                <b style="color:{CATEGORY_COLORS[cat]};font-size:16px;">ประเภทวัสดุ: ประเภท {cat}</b><br>
                <small style="color:#6b7280;">{class_result['detail']}</small>
            </div>""", unsafe_allow_html=True)

        st.markdown("#### Dose Rate ที่แหล่งกำเนิด")
        dose_ref     = st.number_input("Dose rate อ้างอิง (µSv/h)",
                                        min_value=0.0, value=100.0, format="%.2f")
        ref_distance = st.number_input("ที่ระยะ (เมตร)", min_value=0.01,
                                        value=1.0, format="%.2f")
        workload_h   = st.number_input("ชั่วโมงใช้งาน/สัปดาห์ (h/week)",
                                        min_value=0.0, value=8.0)

    st.session_state["src"] = {
        "isotope": isotope, "activity": activity, "unit": unit,
        "activity_tbq": activity_tbq, "use_type": use_type,
        "class_result": class_result,
        "dose_ref": dose_ref, "ref_distance": ref_distance,
        "workload_h": workload_h,
        "energy_group": ISOTOPE_ENERGY_GROUP.get(isotope, "medium"),
    }

# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.subheader("คำนวณกำบังรังสีและกำหนดพื้นที่ (Field use / ใช้งานทั่วไป)")
    st.caption("อ้างอิงสมการที่ 5, 7 จากเอกสาร ปส. และ Inverse square law")

    src = st.session_state.get("src", {})
    dose_ref_t2  = src.get("dose_ref", 100.0)
    ref_dist_t2  = src.get("ref_distance", 1.0)
    workload_h   = src.get("workload_h", 8.0)
    eg_auto      = src.get("energy_group", "medium")

    col_l, col_r = st.columns(2)

    with col_l:
        st.markdown("#### 📏 Inverse Square Law  `I₁d₁² = I₂d₂²`")
        calc_dist = st.number_input("ระยะที่ต้องการคำนวณ (m)",
                                     min_value=0.01, value=2.0, format="%.2f")
        dose_at_dist = dose_ref_t2 * (ref_dist_t2 / calc_dist) ** 2
        dose_annual  = dose_at_dist * workload_h * 52 / 1000  # mSv/yr

        st.markdown(f"""
        <div style="background:#eff6ff;border-left:4px solid #2d7dd2;
                    border-radius:0 8px 8px 0;padding:14px 16px;margin:10px 0;">
            <p style="margin:0;font-size:13px;color:#374151;">
                Dose rate ที่ {calc_dist} m<br>
                <b style="font-size:24px;color:#1a4a7a;">{dose_at_dist:.2f} µSv/h</b>
            </p>
            <p style="margin:6px 0 0;font-size:12px;color:#6b7280;">
                ≈ {dose_annual:.2f} mSv/year (workload {workload_h} h/week × 52 สัปดาห์)
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("#### 🚧 กำหนดพื้นที่ — คำนวณรัศมี")
        st.caption("ขอบเขตที่กั้นบริเวณรังสีต้องมีระดับรังสีไม่เกิน 25 µSv/h (ตาม ปส.)")
        area_defs = [
            ("Controlled Area",          6.0,  25.0,  "#dc2626", "#fef2f2"),
            ("Supervised Area",          1.0,  7.5,   "#d97706", "#fffbeb"),
            ("พื้นที่สาธารณะ / ประชาชน", 1.0,  2.5,   "#16a34a", "#f0fdf4"),
        ]
        area_results = []
        for area_name, limit_msv, limit_usvh, color, bg in area_defs:
            if workload_h > 0:
                r = ref_dist_t2 * math.sqrt(
                    dose_ref_t2 * workload_h * 52 / 1000 / limit_msv)
                r_instant = ref_dist_t2 * math.sqrt(dose_ref_t2 / limit_usvh)
            else:
                r = r_instant = 0
            area_results.append((area_name, limit_msv, r, r_instant, color, bg))
            st.markdown(f"""
            <div style="background:{bg};border:1.5px solid {color};
                        border-radius:8px;padding:10px 14px;margin:6px 0;">
                <b style="color:{color};">{area_name}</b>
                <span style="font-size:12px;color:#6b7280;"> ≤{limit_msv} mSv/yr | ≤{limit_usvh} µSv/h</span><br>
                <span style="font-size:13px;">รัศมีตาม annual dose: <b>{r:.2f} m</b>
                &nbsp;|&nbsp; รัศมีตาม instant dose rate: <b>{r_instant:.2f} m</b></span>
            </div>
            """, unsafe_allow_html=True)
        st.session_state["area_results"] = area_results

    with col_r:
        st.markdown("#### 🧱 คำนวณความหนากำบัง  `I₀/I = 2^(nHVL)`")
        st.caption("สมการที่ 7 จากเอกสาร ปส.")
        eg_opts = list(ENERGY_GROUP_LABEL.keys())
        eg_sel  = st.selectbox("กลุ่มพลังงานโฟตอน", eg_opts,
                                format_func=lambda x: ENERGY_GROUP_LABEL[x],
                                index=eg_opts.index(eg_auto))
        material = st.selectbox("วัสดุกำบัง", MATERIALS)
        limit_area = st.selectbox("Design limit (พื้นที่ปลายทาง)",
                                   list(DOSE_LIMITS_WEEK.keys()))
        limit_usvh_design = DOSE_LIMITS_WEEK[limit_area] / (workload_h if workload_h > 0 else 1)

        hvl = HVL_DATA.get((eg_sel, material))
        tvl = hvl_to_tvl(hvl) if hvl else None

        if hvl and dose_at_dist > 0:
            if dose_at_dist <= limit_usvh_design:
                st.success(f"✅ Dose rate {dose_at_dist:.2f} µSv/h ≤ Design limit "
                           f"{limit_usvh_design:.2f} µSv/h — ไม่ต้องกำบังเพิ่ม")
                n_hvl = n_tvl = thickness_cm = 0.0
            else:
                ratio    = dose_at_dist / limit_usvh_design
                n_hvl    = math.log2(ratio)
                n_tvl    = math.log10(ratio)   # No.of TVL = log₁₀(1/B) = log₁₀(ratio)
                thickness_cm = n_hvl * hvl     # ใช้ HVL ในการคำนวณ

                st.markdown(f"""
                <div style="background:#f0fdf4;border:2px solid #16a34a;
                            border-radius:10px;padding:16px;margin-top:8px;">
                    <p style="margin:0 0 4px;font-size:13px;color:#374151;">ความหนา{material}ที่ต้องการ</p>
                    <p style="margin:0;font-size:30px;font-weight:700;color:#166534;">
                        {thickness_cm:.1f} cm</p>
                    <hr style="border-color:#16a34444;margin:10px 0;">
                    <table style="font-size:12px;color:#374151;width:100%;border-collapse:collapse;">
                        <tr><td>Dose rate (ก่อนกำบัง)</td>
                            <td><b>{dose_at_dist:.2f} µSv/h</b></td></tr>
                        <tr><td>Design limit</td>
                            <td><b>{limit_usvh_design:.2f} µSv/h</b>
                            ({DOSE_LIMITS_WEEK[limit_area]} µSv/week)</td></tr>
                        <tr><td>อัตราส่วน I₀/I</td>
                            <td><b>{ratio:.2f}×</b></td></tr>
                        <tr><td>จำนวน HVL ที่ต้องการ</td>
                            <td><b>{n_hvl:.2f} HVL</b> (HVL={hvl} cm)</td></tr>
                        <tr><td>จำนวน TVL ที่ต้องการ</td>
                            <td><b>{n_tvl:.2f} TVL</b> (TVL≈{tvl:.1f} cm)</td></tr>
                    </table>
                    <p style="margin:8px 0 0;font-size:11px;color:#6b7280;">
                        สูตร: n_HVL = log₂(I₀/I) | thickness = n_HVL × HVL<br>
                        อ้างอิง: IAEA SRS No.47 / NCRP No.151
                    </p>
                </div>
                """, unsafe_allow_html=True)

            st.session_state["shielding"] = {
                "material": material, "thickness_cm": thickness_cm,
                "n_hvl": n_hvl, "n_tvl": n_tvl, "hvl": hvl, "tvl": tvl,
                "energy_group": eg_sel,
                "design_limit_usvh": limit_usvh_design,
                "dose_at_dist": dose_at_dist, "calc_dist": calc_dist,
                "limit_sel": limit_area,
            }

        # Reference table
        with st.expander("📋 ตาราง HVL/TVL อ้างอิง (NCRP 151 / IAEA SRS No.47)"):
            import pandas as pd
            rows = []
            for mat in MATERIALS:
                row = {"วัสดุ": mat}
                for eg in eg_opts:
                    h = HVL_DATA.get((eg, mat), None)
                    t = hvl_to_tvl(h) if h else None
                    row[eg[:6]+"… HVL"] = f"{h} cm" if h else "—"
                    row[eg[:6]+"… TVL"] = f"{t:.1f} cm" if t else "—"
                rows.append(row)
            st.dataframe(pd.DataFrame(rows), hide_index=True, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
with tab3:
    st.subheader("คำนวณผนังห้องฉายรังสี (NCRP/IAEA Method)")
    st.caption(
        "สมการที่ 1–3 จากเอกสาร ปส. อ้างอิง IAEA SRS No.47 | "
        "Dose = Workload × Dose rate × T × U (สมการที่ 4)"
    )

    mode = st.radio("เลือกวิธีคำนวณ", [
        "กรณีที่ 1: คำนวณผนังห้องฉาย (B = Pd²/WUT → No.TVL)",
        "กรณีที่ 2: ตรวจวัด Dose rate แล้วประเมิน (Dose = W × DR × T × U)",
        "กรณีที่ 3: ภาคสนาม (คำนวณรัศมีกั้นพื้นที่ + ความหนากำบัง)",
    ], horizontal=False)

    if "กรณีที่ 1" in mode:
        st.markdown("### 📐 กรณีที่ 1: คำนวณความหนาผนังห้องฉาย")
        st.markdown("""
        **สมการ:**
        - W = RAKR × A × t × n  *(Workload)*
        - B = Pd² / (W × U × T)  *(Required attenuation)*
        - No. of TVL = log₁₀(1/B)  *(จำนวน TVL ที่ต้องการ)*
        """)

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**ข้อมูล Workload**")
            src_rakr = st.selectbox("แหล่งกัมมันตรังสี (RAKR)",
                                     list(RAKR.keys()))
            act_mbq  = st.number_input("Activity (MBq)", min_value=0.0,
                                        value=518000.0, format="%.0f",
                                        help="Ir-192 HDR = 14 Ci ≈ 518,000 MBq")
            t_treat  = st.number_input("เวลารักษาผู้ป่วยต่อราย (h)",
                                        min_value=0.0, value=0.2, format="%.3f",
                                        help="เฉลี่ย 10-15 นาที = 0.167-0.25 h")
            n_pts    = st.number_input("จำนวนผู้ป่วย/สัปดาห์", min_value=0,
                                        value=10, step=1)
            W = RAKR[src_rakr] * act_mbq * t_treat * n_pts
            st.info(f"**Workload W = {W:.2f} µGy·m²/week**")

        with c2:
            st.markdown("**ข้อมูลผนัง**")
            d_wall   = st.number_input("ระยะจากแหล่งถึงผนัง (m)",
                                        min_value=0.1, value=3.0, format="%.1f")
            area_type = st.selectbox("ประเภทพื้นที่โดยรอบผนัง",
                                      ["ผู้ปฏิบัติงาน (P=400 µSv/week)",
                                       "ประชาชนทั่วไป (P=20 µSv/week)"])
            P = 400.0 if "ปฏิบัติงาน" in area_type else 20.0

            T_sel = st.selectbox("Occupancy factor (T)", list(OCCUPANCY_FACTORS.keys()))
            T = OCCUPANCY_FACTORS[T_sel]
            U_sel = st.selectbox("Use factor (U)", list(USE_FACTORS.keys()))
            U = USE_FACTORS[U_sel]

            mat_wall  = st.selectbox("วัสดุผนัง", MATERIALS, key="mat_wall")
            eg_wall   = st.selectbox("กลุ่มพลังงาน", list(ENERGY_GROUP_LABEL.keys()),
                                      format_func=lambda x: ENERGY_GROUP_LABEL[x],
                                      key="eg_wall")

        if st.button("🧮 คำนวณผนังห้องฉาย", type="primary"):
            if W > 0 and U > 0:
                B = (P * d_wall**2) / (W * U * T)
                n_tvl_req = math.log10(1 / B) if B < 1 else 0
                tvl_mat = hvl_to_tvl(HVL_DATA.get((eg_wall, mat_wall), 1))
                thickness = n_tvl_req * tvl_mat

                color = "#dc2626" if B < 1 else "#16a34a"
                bg    = "#fef2f2" if B < 1 else "#f0fdf4"
                st.markdown(f"""
                <div style="background:{bg};border:2px solid {color};
                            border-radius:12px;padding:18px;margin-top:12px;">
                    <p style="margin:0 0 4px;font-size:13px;">ความหนา{mat_wall}ที่ต้องการ</p>
                    <p style="margin:0;font-size:30px;font-weight:700;color:{color};">
                        {thickness:.1f} cm</p>
                    <hr style="border-color:{color}44;margin:10px 0;">
                    <table style="font-size:12px;color:#374151;width:100%;">
                        <tr><td>Workload (W)</td><td><b>{W:.2f} µGy·m²/week</b></td></tr>
                        <tr><td>Dose limit (P)</td><td><b>{P} µSv/week</b></td></tr>
                        <tr><td>ระยะ (d)</td><td><b>{d_wall} m</b></td></tr>
                        <tr><td>T × U</td><td><b>{T} × {U} = {T*U:.3f}</b></td></tr>
                        <tr><td>Required attenuation (B)</td>
                            <td><b>{B:.4e}</b></td></tr>
                        <tr><td>No. of TVL = log₁₀(1/B)</td>
                            <td><b>{n_tvl_req:.2f} TVL</b></td></tr>
                        <tr><td>TVL ของ{mat_wall}</td>
                            <td><b>≈{tvl_mat:.1f} cm</b></td></tr>
                    </table>
                </div>
                """, unsafe_allow_html=True)
                st.session_state["wall_calc"] = {
                    "W": W, "B": B, "n_tvl": n_tvl_req,
                    "thickness": thickness, "material": mat_wall,
                    "d": d_wall, "P": P, "T": T, "U": U,
                }
            else:
                st.warning("กรุณาตรวจสอบค่า U และ Workload ให้มากกว่า 0")

    elif "กรณีที่ 2" in mode:
        st.markdown("### 📊 กรณีที่ 2: ประเมินจากการตรวจวัด Dose rate")
        st.markdown("**Dose = Workload × Dose rate × Occupancy factor (T) × Use factor (U)**")
        st.caption("สมการที่ 4 จากเอกสาร ปส.")

        c1, c2 = st.columns(2)
        with c1:
            dr_measured = st.number_input("Dose rate ที่วัดได้ (µSv/h)",
                                           min_value=0.0, value=20.0, format="%.2f")
            wl_h_week   = st.number_input("Workload (h/week)",
                                           min_value=0.0, value=1.94, format="%.3f",
                                           help="เวลาฉายรังสีต่อสัปดาห์ (h)")
            T2 = st.selectbox("Occupancy factor (T)", list(OCCUPANCY_FACTORS.keys()),
                               key="T2")
            U2 = st.selectbox("Use factor (U)", list(USE_FACTORS.keys()), key="U2")
            T2v = OCCUPANCY_FACTORS[T2]
            U2v = USE_FACTORS[U2]

        with c2:
            dose_week = wl_h_week * dr_measured * T2v * U2v
            dose_year = dose_week * 52 / 1000   # mSv/yr
            limit_week_sel = st.selectbox("Dose limit ที่ใช้เปรียบเทียบ",
                                           list(DOSE_LIMITS_WEEK.keys()), key="lim2")
            limit_val = DOSE_LIMITS_WEEK[limit_week_sel]
            safe = dose_week <= limit_val

            color = "#16a34a" if safe else "#dc2626"
            bg    = "#f0fdf4" if safe else "#fef2f2"
            status = "✅ ปลอดภัย" if safe else "⚠️ เกิน limit!"
            st.markdown(f"""
            <div style="background:{bg};border:2px solid {color};
                        border-radius:10px;padding:16px;">
                <p style="margin:0 0 4px;font-size:13px;">Dose ที่ประเมินได้</p>
                <p style="margin:0;font-size:26px;font-weight:700;color:{color};">
                    {dose_week:.2f} µSv/week</p>
                <p style="margin:4px 0;font-size:14px;color:{color};">
                    = {dose_year:.3f} mSv/year &nbsp; {status}</p>
                <hr style="border-color:{color}44;margin:10px 0;">
                <p style="font-size:12px;color:#374151;margin:0;">
                    Dose limit: <b>{limit_val} µSv/week</b><br>
                    W={wl_h_week} h/week | DR={dr_measured} µSv/h | T={T2v} | U={U2v}
                </p>
            </div>
            """, unsafe_allow_html=True)

    else:  # กรณีที่ 3 ภาคสนาม
        st.markdown("### 🏗️ กรณีที่ 3: ภาคสนาม — คำนวณรัศมีกั้นพื้นที่และความหนากำบัง")
        st.caption("ขอบเขตที่กั้นต้องมีระดับรังสีไม่เกิน 25 µSv/h (ตาม ปส.)")

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Part A: คำนวณรัศมีกั้นพื้นที่**")
            I1 = st.number_input("Dose rate อ้างอิง I₁ (µSv/h)",
                                  min_value=0.0, value=2560.0, format="%.1f",
                                  help="เช่น 2.56 mSv/h = 2560 µSv/h")
            d1 = st.number_input("ที่ระยะ d₁ (m)", min_value=0.01, value=1.0)
            I2_target = st.number_input("Dose rate เป้าหมาย I₂ (µSv/h)",
                                         min_value=0.1, value=25.0,
                                         help="ขอบเขตพื้นที่ปฏิบัติงาน = 25 µSv/h")
            if I2_target > 0:
                d2 = d1 * math.sqrt(I1 / I2_target)
                color = "#1a4a7a"
                st.markdown(f"""
                <div style="background:#eff6ff;border:2px solid #2d7dd2;
                            border-radius:10px;padding:14px;margin-top:10px;">
                    <p style="margin:0;font-size:13px;">รัศมีกั้นพื้นที่</p>
                    <p style="margin:0;font-size:28px;font-weight:700;color:#1a4a7a;">
                        d₂ = {d2:.1f} เมตร</p>
                    <p style="font-size:12px;color:#6b7280;margin:4px 0 0;">
                        สูตร: d₂ = d₁ × √(I₁/I₂) = {d1}×√({I1}/{I2_target})<br>
                        สมการที่ 5: I₁d₁² = I₂d₂²
                    </p>
                </div>
                """, unsafe_allow_html=True)

        with c2:
            st.markdown("**Part B: คำนวณความหนากำบัง (กรณีพื้นที่คับแคบ)**")
            I0_f = st.number_input("Dose rate เริ่มต้น I₀ (µSv/h)",
                                    min_value=0.0, value=2560.0, format="%.1f")
            I_f  = st.number_input("Dose rate เป้าหมาย I (µSv/h)",
                                    min_value=0.1, value=25.0)
            eg_f = st.selectbox("กลุ่มพลังงาน", list(ENERGY_GROUP_LABEL.keys()),
                                  format_func=lambda x: ENERGY_GROUP_LABEL[x], key="eg_f")
            mat_f = st.selectbox("วัสดุกำบัง", MATERIALS, key="mat_f")

            hvl_f = HVL_DATA.get((eg_f, mat_f))
            if hvl_f and I_f > 0 and I0_f > 0:
                n_hvl_f = math.log2(I0_f / I_f)
                thick_f = n_hvl_f * hvl_f
                st.markdown(f"""
                <div style="background:#f0fdf4;border:2px solid #16a34a;
                            border-radius:10px;padding:14px;margin-top:10px;">
                    <p style="margin:0;font-size:13px;">ความหนา{mat_f}ที่ต้องการ</p>
                    <p style="margin:0;font-size:28px;font-weight:700;color:#166534;">
                        {thick_f:.1f} cm ({thick_f*10:.1f} mm)</p>
                    <hr style="border-color:#16a34444;margin:8px 0;">
                    <p style="font-size:12px;color:#374151;margin:0;">
                        I₀/I = {I0_f}/{I_f} = {I0_f/I_f:.1f}<br>
                        n_HVL = log₂({I0_f/I_f:.1f}) = {n_hvl_f:.2f}<br>
                        HVL({mat_f}) = {hvl_f} cm<br>
                        <i>สมการที่ 7: I₀/I = 2^(nHVL)</i>
                    </p>
                </div>
                """, unsafe_allow_html=True)

        # D = R × T (เวลาทำงานสูงสุด)
        st.divider()
        st.markdown("#### ⏱️ คำนวณเวลาทำงานสูงสุด  `D = R × T`")
        st.caption("สมการที่ 8 จากเอกสาร ปส. | D=10 µSv/h (สำหรับผู้ปฏิบัติงาน)")
        c3, c4 = st.columns(2)
        with c3:
            R_work = st.number_input("Dose rate ณ จุดปฏิบัติงาน (µSv/h)",
                                      min_value=0.0, value=25.0, format="%.2f")
            D_limit = st.selectbox("Dose limit",
                                    ["10 µSv/h (ผู้ปฏิบัติงาน, 8h/day)",
                                     "400 µSv/week (ผู้ปฏิบัติงาน)",
                                     "20 µSv/week (ประชาชน)"])
            D_val = 10 if "10 µSv/h" in D_limit else (400 if "400" in D_limit else 20)
            unit_d = "h" if "10 µSv/h" in D_limit else "week"
        with c4:
            if R_work > 0:
                T_max = D_val / R_work
                st.markdown(f"""
                <div style="background:#eff6ff;border:2px solid #2d7dd2;
                            border-radius:10px;padding:14px;margin-top:8px;">
                    <p style="margin:0;font-size:13px;">เวลาทำงานสูงสุดที่อนุญาต</p>
                    <p style="margin:0;font-size:28px;font-weight:700;color:#1a4a7a;">
                        {T_max:.2f} {unit_d}</p>
                    <p style="font-size:12px;color:#6b7280;margin:4px 0 0;">
                        T = D/R = {D_val}/{R_work:.2f} = {T_max:.2f} {unit_d}<br>
                        (D limit = {D_val} µSv/{unit_d})
                    </p>
                </div>
                """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
with tab4:
    st.subheader("คำนวณ Dose ผู้ปฏิบัติงาน")
    src4 = st.session_state.get("src", {})

    if "workers" not in st.session_state:
        st.session_state.workers = [
            {"name": "RSO",             "distance": 2.0, "hours_week": 2.0},
            {"name": "ผู้ปฏิบัติงาน 1", "distance": 1.5, "hours_week": 8.0},
            {"name": "ผู้ปฏิบัติงาน 2", "distance": 3.0, "hours_week": 4.0},
        ]

    workers_updated = []
    for i, w in enumerate(st.session_state.workers):
        wc1, wc2, wc3, wc4 = st.columns([2, 1.5, 1.5, 0.5])
        name  = wc1.text_input(f"ชื่อ #{i+1}", value=w["name"], key=f"wn_{i}")
        dist  = wc2.number_input(f"ระยะทาง #{i+1} (m)", min_value=0.01,
                                   value=float(w["distance"]), format="%.2f", key=f"wd_{i}")
        hours = wc3.number_input(f"ชั่วโมง/สัปดาห์ #{i+1}", min_value=0.0,
                                   value=float(w["hours_week"]), format="%.1f", key=f"wh_{i}")
        wc4.markdown("<br>", unsafe_allow_html=True)
        if not wc4.button("🗑️", key=f"wdel_{i}"):
            workers_updated.append({"name": name, "distance": dist, "hours_week": hours})
    st.session_state.workers = workers_updated

    if st.button("➕ เพิ่มผู้ปฏิบัติงาน"):
        st.session_state.workers.append({
            "name": f"ผู้ปฏิบัติงาน {len(st.session_state.workers)+1}",
            "distance": 2.0, "hours_week": 4.0})
        st.rerun()

    st.divider()
    if st.button("🧮 คำนวณ Dose ทั้งหมด", type="primary"):
        dr_base  = src4.get("dose_ref", 100.0)
        rd_base  = src4.get("ref_distance", 1.0)
        results  = []
        import pandas as pd
        for w in st.session_state.workers:
            dr = dr_base * (rd_base / w["distance"]) ** 2
            annual = dr * w["hours_week"] * 52 / 1000
            pct    = annual / 20.0 * 100
            t_max  = 10.0 / dr if dr > 0 else float("inf")  # D=R×T → T=D/R
            status = "✅ ปลอดภัย" if annual <= 20 else "⚠️ เกิน 20 mSv/yr!"
            results.append({
                "ชื่อ": w["name"],
                "ระยะ (m)": w["distance"],
                "Dose rate (µSv/h)": round(dr, 2),
                "Annual Dose (mSv/yr)": round(annual, 3),
                "% ของ 20 mSv/yr": round(pct, 1),
                "T_max (h/วัน)": f"{t_max:.1f}" if t_max < 1000 else "ไม่จำกัด",
                "สถานะ": status,
            })
        df = pd.DataFrame(results)
        st.dataframe(df, hide_index=True, use_container_width=True,
                     column_config={
                         "% ของ 20 mSv/yr": st.column_config.ProgressColumn(
                             min_value=0, max_value=100, format="%.1f%%"),
                     })
        st.session_state["worker_results"] = results
        over = [r for r in results if "เกิน" in r["สถานะ"]]
        if over:
            st.error(f"⚠️ มี {len(over)} คนที่ได้รับ Dose เกิน 20 mSv/yr — "
                     f"กรุณาเพิ่มกำบัง ปรับระยะทาง หรือลดชั่วโมงทำงาน")
        else:
            st.success("✅ ผู้ปฏิบัติงานทุกคนอยู่ภายใต้ 20 mSv/yr ค่ะ")

# ══════════════════════════════════════════════════════════════════════════════
with tab5:
    st.subheader("AI ร่างแผนป้องกันอันตรายจากรังสี")

    src5    = st.session_state.get("src", {})
    shield5 = st.session_state.get("shielding", {})
    areas5  = st.session_state.get("area_results", [])
    workers5= st.session_state.get("worker_results", [])
    wall5   = st.session_state.get("wall_calc", {})

    if not anthropic_key:
        st.warning("⚠️ กรุณาใส่ Anthropic API Key ใน sidebar")
    else:
        if st.button("✨ ร่างแผนป้องกันอันตรายจากรังสี", type="primary"):
            area_text = "\n".join(
                [f"  - {a[0]}: รัศมี annual dose ≥ {a[2]:.2f} m, instant dose ≥ {a[3]:.2f} m"
                 for a in areas5]) if areas5 else "ยังไม่ได้คำนวณ"

            worker_text = "\n".join(
                [f"  - {w['ชื่อ']}: {w['Annual Dose (mSv/yr)']} mSv/yr ({w['สถานะ']})"
                 for w in workers5]) if workers5 else "ยังไม่ได้คำนวณ"

            shield_text = (
                f"ความหนา{shield5.get('material','—')}: {shield5.get('thickness_cm',0):.1f} cm "
                f"({shield5.get('n_hvl',0):.2f} HVL, HVL={shield5.get('hvl','—')} cm, "
                f"TVL≈{shield5.get('tvl','—')} cm)"
            ) if shield5 else "ยังไม่ได้คำนวณ"

            wall_text = (
                f"ผนังห้องฉาย {wall5.get('material','—')}: {wall5.get('thickness',0):.1f} cm "
                f"({wall5.get('n_tvl',0):.2f} TVL, B={wall5.get('B',0):.2e})"
            ) if wall5 else "ยังไม่ได้คำนวณ"

            prompt = f"""คุณเป็นผู้เชี่ยวชาญด้านความปลอดภัยทางรังสีและกฎหมายพลังงานนิวเคลียร์ของไทย
จงร่างแผนป้องกันอันตรายจากรังสีที่สมบูรณ์ตามข้อกำหนดของสำนักงานปรมาณูเพื่อสันติ (ปส.)
อ้างอิง: กฎกระทรวงความปลอดภัยทางรังสี พ.ศ. 2561 | IAEA SRS No.47 | NCRP No.151

ข้อมูลองค์กร:
- องค์กร: {org_name} | ใบอนุญาต: {license_no}
- RSO: {rso_name} โทร. {rso_phone} | วันที่: {plan_date}

ข้อมูลวัสดุกัมมันตรังสี:
- ชนิด: {src5.get('isotope','—')} | Activity: {src5.get('activity','—')} {src5.get('unit','—')}
  ({src5.get('activity_tbq',0):.3e} TBq)
- ประเภทวัสดุ: ประเภท {src5.get('class_result',{}).get('category','—')}
- Dose rate: {src5.get('dose_ref',0)} µSv/h ที่ระยะ {src5.get('ref_distance',1)} m
- Workload: {src5.get('workload_h',0)} h/week

ผลการคำนวณ:
- การกำหนดพื้นที่:
{area_text}
- กำบังรังสี (ภาคสนาม/ทั่วไป): {shield_text}
- กำบังรังสี (ผนังห้องฉาย): {wall_text}
- Dose ผู้ปฏิบัติงาน:
{worker_text}

จงร่างแผนป้องกันอันตรายจากรังสีภาษาไทยอย่างเป็นทางการ ประกอบด้วยหัวข้อตามที่กฎหมายกำหนด:
1. แผนผังสายการบังคับบัญชาด้านความปลอดภัยทางรังสี
2. การจัดแบ่งพื้นที่ (Controlled Area / Supervised Area) พร้อมรัศมีที่คำนวณได้
3. กฎระเบียบและมาตรการความปลอดภัยทางรังสี (หลัก TID: Time, Distance, Shielding)
4. แผนการตรวจวัดรังสีประจำพื้นที่และบุคคล (OSL/TLD) พร้อมความถี่
5. เครื่องมืออุปกรณ์ป้องกันอันตรายจากรังสี
6. ข้อกำหนดกำบังรังสี (ผลการคำนวณจาก NCRP/IAEA)
7. แผนปฏิบัติกรณีเกิดเหตุฉุกเฉินทางรังสี (อ้างอิงคู่มือ ปส.)
8. แผนการฝึกอบรมบุคลากร
9. ระบบบันทึกและรายงานด้านความปลอดภัย"""

            with st.spinner("AI กำลังร่างแผน..."):
                try:
                    resp = requests.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={"Content-Type": "application/json",
                                 "x-api-key": anthropic_key,
                                 "anthropic-version": "2023-06-01"},
                        json={"model": "claude-sonnet-4-20250514",
                              "max_tokens": 2000,
                              "messages": [{"role": "user", "content": prompt}]},
                        timeout=60,
                    )
                    data = resp.json()
                    draft = data["content"][0]["text"]
                    st.session_state["protection_draft"] = draft
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {e}")

    if "protection_draft" in st.session_state:
        st.warning("⚠️ กรุณาตรวจสอบและแก้ไขก่อนใช้งานจริง")
        edited = st.text_area("แผนป้องกันอันตรายจากรังสี (แก้ไขได้โดยตรง)",
                               value=st.session_state["protection_draft"], height=500)
        st.session_state["protection_draft_edited"] = edited

        st.divider()
        if st.button("📥 Export Word (.docx)", type="primary"):
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io, pandas as pd

                doc = Document()
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = title.add_run("แผนป้องกันอันตรายจากรังสี")
                r.bold = True; r.font.size = Pt(18)
                r.font.color.rgb = RGBColor(0x1a, 0x4a, 0x7a)
                doc.add_paragraph(
                    "อ้างอิง: กฎกระทรวงความปลอดภัยทางรังสี พ.ศ. 2561 | "
                    "IAEA SRS No.47 | NCRP No.151 | เอกสาร ปส."
                ).alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

                def add_table(doc, rows_data):
                    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
                    for lbl, val in rows_data:
                        row = t.add_row().cells
                        row[0].text = lbl; row[1].text = str(val)

                doc.add_heading("1. ข้อมูลองค์กรและผู้รับผิดชอบ", level=1)
                add_table(doc, [
                    ("ชื่อองค์กร", org_name), ("เลขที่ใบอนุญาต", license_no),
                    ("RSO", rso_name), ("โทรศัพท์ RSO", rso_phone),
                    ("วันที่จัดทำ", str(plan_date)),
                ])
                doc.add_paragraph()

                doc.add_heading("2. ข้อมูลวัสดุกัมมันตรังสี", level=1)
                cr = src5.get("class_result", {})
                add_table(doc, [
                    ("ชนิดนิวไคลด์", src5.get("isotope", "—")),
                    ("ค่ากัมมันตภาพ",
                     f"{src5.get('activity','—')} {src5.get('unit','—')} "
                     f"({src5.get('activity_tbq',0):.3e} TBq)"),
                    ("ประเภทวัสดุ", f"ประเภท {cr.get('category','—')}"),
                    ("Dose rate", f"{src5.get('dose_ref',0)} µSv/h ที่ {src5.get('ref_distance',1)} m"),
                    ("Workload", f"{src5.get('workload_h',0)} h/week"),
                ])
                doc.add_paragraph()

                if areas5:
                    doc.add_heading("3. การกำหนดพื้นที่ควบคุมและตรวจตรา", level=1)
                    t = doc.add_table(rows=0, cols=4); t.style = "Table Grid"
                    hdr = t.add_row().cells
                    hdr[0].text = "พื้นที่"; hdr[1].text = "Limit (mSv/yr)"
                    hdr[2].text = "รัศมี Annual (m)"; hdr[3].text = "รัศมี Instant (m)"
                    for a in areas5:
                        row = t.add_row().cells
                        row[0].text = a[0]; row[1].text = str(a[1])
                        row[2].text = f"≥ {a[2]:.2f}"; row[3].text = f"≥ {a[3]:.2f}"
                    doc.add_paragraph()

                if shield5:
                    doc.add_heading("4. ผลการคำนวณกำบังรังสี", level=1)
                    add_table(doc, [
                        ("วัสดุกำบัง", shield5.get("material","—")),
                        ("ความหนาที่ต้องการ", f"{shield5.get('thickness_cm',0):.1f} cm"),
                        ("จำนวน HVL", f"{shield5.get('n_hvl',0):.2f}"),
                        ("HVL", f"{shield5.get('hvl','—')} cm"),
                        ("TVL", f"≈{shield5.get('tvl','—')} cm"),
                    ])
                    doc.add_paragraph()

                if wall5:
                    doc.add_heading("5. ผลการคำนวณผนังห้องฉายรังสี (NCRP/IAEA)", level=1)
                    add_table(doc, [
                        ("วัสดุผนัง", wall5.get("material","—")),
                        ("ความหนาที่ต้องการ", f"{wall5.get('thickness',0):.1f} cm"),
                        ("Required attenuation (B)", f"{wall5.get('B',0):.2e}"),
                        ("จำนวน TVL", f"{wall5.get('n_tvl',0):.2f}"),
                    ])
                    doc.add_paragraph()

                if workers5:
                    doc.add_heading("6. Dose ผู้ปฏิบัติงาน", level=1)
                    t2 = doc.add_table(rows=0, cols=5); t2.style = "Table Grid"
                    hdr2 = t2.add_row().cells
                    hdr2[0].text = "ชื่อ"; hdr2[1].text = "ระยะ (m)"
                    hdr2[2].text = "Annual Dose (mSv/yr)"
                    hdr2[3].text = "T_max (h/วัน)"; hdr2[4].text = "สถานะ"
                    for w in workers5:
                        row2 = t2.add_row().cells
                        row2[0].text = w["ชื่อ"]; row2[1].text = str(w["ระยะ (m)"])
                        row2[2].text = str(w["Annual Dose (mSv/yr)"])
                        row2[3].text = str(w.get("T_max (h/วัน)","—"))
                        row2[4].text = w["สถานะ"]
                    doc.add_paragraph()

                doc.add_heading("7. รายละเอียดแผนป้องกันอันตราย (AI Draft)", level=1)
                draft_f = st.session_state.get("protection_draft_edited",
                          st.session_state.get("protection_draft",""))
                for line in draft_f.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

                doc.add_paragraph()
                doc.add_paragraph("ลงนาม: ________________________________  วันที่: _______________")
                doc.add_paragraph("(ผู้รับใบอนุญาต / เจ้าหน้าที่ความปลอดภัยทางรังสี)")

                buf = io.BytesIO(); doc.save(buf); buf.seek(0)
                st.download_button(
                    label="⬇️ ดาวน์โหลด radiation_protection_plan.docx",
                    data=buf, file_name="radiation_protection_plan.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("✅ ไฟล์ Word พร้อมดาวน์โหลดแล้วค่ะ")
            except ImportError:
                st.error("กรุณาติดตั้ง python-docx: `pip install python-docx`")
            except Exception as e:
                st.error(f"เกิดข้อผิดพลาด: {e}")
