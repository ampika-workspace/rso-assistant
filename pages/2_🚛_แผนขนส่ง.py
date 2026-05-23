"""
🚛 Page 2 — แผนการขนส่งวัสดุกัมมันตรังสี
คำนวณ TI, Category, Package Type, UN Number, Label
และ AI ร่างแผนพร้อม export Word
อ้างอิง IAEA SSR-6 (Rev.1) 2018
"""
import streamlit as st
import math, sys, os, requests
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from utils.data import (D_VALUES_TBQ, ISOTOPE_INFO, UNIT_TO_TBQ,
                  classify_material, USE_TYPE_CATEGORY,
                  CATEGORY_COLORS, CATEGORY_BG)

st.set_page_config(page_title="แผนขนส่งวัสดุกัมมันตรังสี", page_icon="🚛", layout="wide")
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@400;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Sarabun', sans-serif !important; }
.result-box { border-radius: 10px; padding: 14px 18px; margin: 6px 0; border: 2px solid; }
</style>
""", unsafe_allow_html=True)

# ── UN Number / Package classification logic (IAEA SSR-6) ─────────────────────
def classify_package(isotope: str, activity_tbq: float, special_form: bool, fissile: bool) -> dict:
    """
    จำแนก Package Type และ UN Number ตาม IAEA SSR-6 Table 1
    
    Logic:
      1. Excepted Package: Activity ≤ exempt limit (ประมาณ A2/1000 สำหรับ non-special)
         → UN 2910 (limited quantity), UN 2911 (instruments)
      2. Type A: Activity ≤ A1 (special form) หรือ ≤ A2 (non-special form)
         → UN 2915 (non-special form), UN 2916 (special form ถ้า Type B)
      3. Type B(U): Activity > A1 หรือ A2
         → UN 2916 (non-fissile), UN 2917 (B(M))
    """
    info = ISOTOPE_INFO.get(isotope, {})
    a1 = info.get("a1")   # TBq — limit สำหรับ special form
    a2 = info.get("a2")   # TBq — limit สำหรับ non-special form

    if not a1 or not a2 or not activity_tbq or activity_tbq <= 0:
        return {"package_type": "—", "un_number": "—", "proper_name": "—",
                "label": "—", "detail": "ข้อมูลไม่เพียงพอ", "color": "#6b7280"}

    # Excepted package limit ≈ A2 / 1000  (para 522, SSR-6)
    # แต่สำหรับความปลอดภัย ใช้เกณฑ์ที่ชัดเจนกว่า:
    # Excepted ถ้า activity ≤ A2/1000 (ของเหลว) หรือ A2/10 (ของแข็ง/เครื่องมือ)
    excepted_limit = a2 / 100  # เกณฑ์ conservative สำหรับ limited quantity

    if activity_tbq <= excepted_limit:
        return {
            "package_type": "Excepted Package",
            "un_number": "UN 2910",
            "proper_name": "Radioactive material, excepted package — limited quantity of material",
            "label": "ไม่ต้องติด Label (ต้องมีเครื่องหมาย UN 2910)",
            "detail": f"Activity ({activity_tbq:.3e} TBq) ≤ Excepted limit ({excepted_limit:.2e} TBq ≈ A₂/100)",
            "color": "#0891b2",
            "badge_color": "#ecfeff",
        }

    # Type A: special form → ≤ A1, non-special form → ≤ A2
    type_a_limit = a1 if special_form else a2

    if activity_tbq <= type_a_limit:
        if special_form:
            return {
                "package_type": "Type A (Special Form)",
                "un_number": "UN 2915",
                "proper_name": "Radioactive material, Type A package, special form [non fissile or fissile-excepted]",
                "label": "Label I-WHITE, II-YELLOW หรือ III-YELLOW (ขึ้นกับ TI)",
                "detail": f"Activity ({activity_tbq:.3e} TBq) ≤ A₁ ({a1} TBq) — Special Form",
                "color": "#16a34a",
                "badge_color": "#f0fdf4",
            }
        else:
            return {
                "package_type": "Type A (Non-Special Form)",
                "un_number": "UN 2915",
                "proper_name": "Radioactive material, Type A package [non-special form, non fissile or fissile-excepted]",
                "label": "Label I-WHITE, II-YELLOW หรือ III-YELLOW (ขึ้นกับ TI)",
                "detail": f"Activity ({activity_tbq:.3e} TBq) ≤ A₂ ({a2} TBq) — Non-Special Form",
                "color": "#16a34a",
                "badge_color": "#f0fdf4",
            }

    # Type B: activity > A1 (special) หรือ > A2 (non-special)
    return {
        "package_type": "Type B(U)",
        "un_number": "UN 2916",
        "proper_name": "Radioactive material, Type B(U) package [non fissile or fissile-excepted]",
        "label": "Label III-YELLOW + Placard",
        "detail": (f"Activity ({activity_tbq:.3e} TBq) > "
                   f"{'A₁' if special_form else 'A₂'} "
                   f"({'%.3g' % a1 if special_form else '%.3g' % a2} TBq) — ต้องใช้ Type B"),
        "color": "#dc2626",
        "badge_color": "#fef2f2",
    }


def classify_label(ti: float) -> dict:
    """จำแนก Label ตาม TI"""
    if ti <= 0.05:
        return {"label": "Label I — WHITE", "color": "#ffffff",
                "border": "#16a34a", "text": "#166534", "category": "Category I-WHITE"}
    elif ti <= 1.0:
        return {"label": "Label II — YELLOW", "color": "#fef08a",
                "border": "#d97706", "text": "#92400e", "category": "Category II-YELLOW"}
    else:
        return {"label": "Label III — YELLOW", "color": "#fde047",
                "border": "#dc2626", "text": "#991b1b", "category": "Category III-YELLOW"}


# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🏢 ข้อมูลองค์กร")
    org_name    = st.text_input("ชื่อองค์กร / ผู้รับใบอนุญาต", placeholder="บริษัท ABC จำกัด")
    license_no  = st.text_input("เลขที่ใบอนุญาต", placeholder="4I0001/63F")
    org_address = st.text_area("ที่อยู่", height=80)
    org_phone   = st.text_input("โทรศัพท์", placeholder="02-xxx-xxxx")
    rso_name    = st.text_input("ชื่อ RSO", placeholder="นายปรมาณู เพื่อสันติ")
    rso_phone   = st.text_input("โทรศัพท์ RSO", placeholder="08x-xxx-xxxx")
    plan_date   = st.date_input("วันที่จัดทำแผน")
    st.markdown("---")
    anthropic_key = st.text_input("🔑 Anthropic API Key", type="password")

st.markdown("# 🚛 แผนการขนส่งวัสดุกัมมันตรังสี")
st.caption("อ้างอิง IAEA SSR-6 (Rev.1) 2018 และสำนักงานปรมาณูเพื่อสันติ")

# ── Tabs ───────────────────────────────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs(["📦 ข้อมูลวัสดุและการขนส่ง", "🤖 AI ร่างแผน", "📄 Export Word"])

isotope_list = list(D_VALUES_TBQ.keys())
use_list = ["— ไม่พบในตาราง / คำนวณ A/D —"] + list(USE_TYPE_CATEGORY.keys())

# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    col_left, col_right = st.columns([1, 1])

    # ── ซ้าย: ข้อมูลวัสดุ ─────────────────────────────────────────────────────
    with col_left:
        st.markdown("#### ☢️ ข้อมูลวัสดุกัมมันตรังสี")

        isotope   = st.selectbox("ชนิดนิวไคลด์", isotope_list,
                                  index=isotope_list.index("Ir-192"))
        c1, c2    = st.columns(2)
        activity  = c1.number_input("ค่ากัมมันตภาพ", min_value=0.0,
                                     value=5.5, format="%.4f")
        unit      = c2.selectbox("หน่วย", list(UNIT_TO_TBQ.keys()),
                                  index=list(UNIT_TO_TBQ.keys()).index("Ci"))
        act_date  = st.date_input("วันที่อ้างอิงกัมมันตภาพ")
        use_type  = st.selectbox("การใช้ประโยชน์", use_list)
        phys_form = st.selectbox("สถานะทางกายภาพ", ["ของแข็ง", "ของเหลว", "ก๊าซ", "Sealed Source"])

        c3, c4 = st.columns(2)
        special_form = c3.checkbox("Special Form", value=True,
                                    help="วัสดุที่ไม่กระจายตัว เช่น Sealed Source — ใช้ A1 เป็นเกณฑ์")
        fissile      = c4.checkbox("Fissile Material", value=False,
                                    help="วัสดุที่แตกตัวได้ เช่น U-235, Pu-239")

        activity_tbq = activity * UNIT_TO_TBQ.get(unit, 1)

        # แสดงข้อมูล isotope
        if isotope in ISOTOPE_INFO:
            info = ISOTOPE_INFO[isotope]
            d_val = D_VALUES_TBQ.get(isotope, "—")
            a1_val = info.get("a1", "—")
            a2_val = info.get("a2", "—")
            st.info(
                f"**{isotope}** | ครึ่งชีวิต: {info['halfLife']} | "
                f"A₁: **{a1_val} TBq** | A₂: **{a2_val} TBq** | D: {d_val} TBq"
            )
            st.caption(f"Activity ที่กรอก = **{activity_tbq:.4e} TBq**")

        # จำแนกประเภทวัสดุ (A/D)
        ut = use_type if use_type != "— ไม่พบในตาราง / คำนวณ A/D —" else ""
        class_result = classify_material(isotope, activity_tbq, ut)
        if class_result.get("category"):
            cat = class_result["category"]
            st.markdown(f"""
            <div style="background:{CATEGORY_BG[cat]}; border:2px solid {CATEGORY_COLORS[cat]};
                        border-radius:8px; padding:10px 14px; margin-top:6px;">
                <b style="color:{CATEGORY_COLORS[cat]};">ประเภทวัสดุ: ประเภท {cat}</b>
                &nbsp;|&nbsp; {class_result['threshold']}<br>
                <small style="color:#6b7280;">{class_result['detail']}</small>
            </div>
            """, unsafe_allow_html=True)

    # ── ขวา: ประเมินบรรจุภัณฑ์ + TI ──────────────────────────────────────────
    with col_right:
        st.markdown("#### 📦 Package Type, UN Number & Label")

        # Package classification
        pkg_result = classify_package(isotope, activity_tbq, special_form, fissile)

        st.markdown(f"""
        <div style="background:{pkg_result.get('badge_color','#f9fafb')};
                    border:2px solid {pkg_result['color']};
                    border-radius:12px; padding:16px; margin-bottom:14px;">
            <div style="display:flex; align-items:center; gap:10px; margin-bottom:10px;">
                <span style="background:{pkg_result['color']}; color:white;
                             padding:5px 14px; border-radius:20px; font-weight:700; font-size:14px;">
                    {pkg_result['package_type']}
                </span>
                <span style="font-weight:700; color:{pkg_result['color']}; font-size:16px;">
                    {pkg_result['un_number']}
                </span>
            </div>
            <p style="margin:0 0 6px; font-size:13px; color:#374151;">
                <b>Proper Shipping Name:</b><br>{pkg_result['proper_name']}
            </p>
            <p style="margin:0 0 4px; font-size:12px; color:#6b7280;">
                📌 {pkg_result['detail']}
            </p>
        </div>
        """, unsafe_allow_html=True)

        # A1 / A2 comparison breakdown
        if isotope in ISOTOPE_INFO:
            info = ISOTOPE_INFO[isotope]
            a1 = info.get("a1", 0)
            a2 = info.get("a2", 0)
            limit_used = a1 if special_form else a2
            ratio = activity_tbq / limit_used if limit_used > 0 else 0

            st.markdown("**เปรียบเทียบ Activity กับ A1/A2:**")
            m1, m2, m3 = st.columns(3)
            m1.metric("Activity (TBq)", f"{activity_tbq:.3e}")
            m2.metric("A₁ (TBq)" if special_form else "A₂ (TBq)",
                      f"{a1 if special_form else a2}")
            m3.metric("สัดส่วน Act/Limit", f"{ratio:.2f}×",
                      delta="เกิน Type A" if ratio > 1 else "อยู่ใน Type A",
                      delta_color="inverse" if ratio > 1 else "normal")

        st.divider()

        # TI & Label
        st.markdown("#### 🏷️ Transport Index (TI) & Label")
        dose_1m = st.number_input("Dose rate ที่ระยะ 1 เมตร (mSv/h)",
                                   min_value=0.0, value=0.5, format="%.4f",
                                   help="วัดจากพื้นผิวภายนอกหีบห่อ ห่างออกมา 1 เมตร")

        if dose_1m >= 0:
            ti = math.ceil(dose_1m * 10) / 10 if dose_1m > 0 else 0
            label_info = classify_label(ti)

            l1, l2 = st.columns(2)
            l1.metric("Transport Index (TI)", ti)

            l2.markdown(f"""
            <div style="background:{label_info['color']};
                        border:3px solid {label_info['border']};
                        border-radius:8px; padding:10px; text-align:center; margin-top:8px;">
                <b style="color:{label_info['text']}; font-size:13px;">
                    {label_info['label']}
                </b><br>
                <small style="color:{label_info['text']};">{label_info['category']}</small>
            </div>
            """, unsafe_allow_html=True)
        else:
            ti, label_info = 0, classify_label(0)

        st.divider()

        # ── Transport info ──────────────────────────────────────────────────────
        st.markdown("#### 🚛 ข้อมูลการขนส่ง")
        transport_mode = st.selectbox("รูปแบบการขนส่ง",
                                       ["ทางถนน", "ทางอากาศ", "ทางเรือ", "ทางรถไฟ"])
        c5, c6 = st.columns(2)
        origin      = c5.text_input("ต้นทาง", placeholder="ชื่อสถานที่")
        destination = c6.text_input("ปลายทาง", placeholder="ชื่อสถานที่")
        c7, c8 = st.columns(2)
        transporter = c7.text_input("บริษัท/ชื่อผู้ขนส่ง")
        trans_phone = c8.text_input("โทรศัพท์ผู้ขนส่ง", placeholder="08x-xxx-xxxx")
        c9, c10 = st.columns(2)
        vehicle     = c9.text_input("ประเภทยานพาหนะ", placeholder="รถกระบะ / รถบรรทุก")
        plate       = c10.text_input("ทะเบียน", placeholder="กก-1234")
        emg1 = st.text_input("ผู้ติดต่อฉุกเฉิน 1 (ชื่อ + โทร.)")
        emg2 = st.text_input("ผู้ติดต่อฉุกเฉิน 2 (สำรอง)")

    # ── Summary box ────────────────────────────────────────────────────────────
    st.divider()
    st.markdown("### 📋 สรุปผลการประเมินบรรจุภัณฑ์และการขนส่ง")
    s1, s2, s3, s4, s5 = st.columns(5)
    s1.markdown(f"**ประเภทวัสดุ**<br>ประเภท {class_result.get('category','—')}",
                unsafe_allow_html=True)
    s2.markdown(f"**Package Type**<br>{pkg_result['package_type']}",
                unsafe_allow_html=True)
    s3.markdown(f"**UN Number**<br>{pkg_result['un_number']}",
                unsafe_allow_html=True)
    s4.markdown(f"**Transport Index**<br>{ti}",
                unsafe_allow_html=True)
    s5.markdown(f"**Label**<br>{label_info['label']}",
                unsafe_allow_html=True)

    st.caption("⚠️ ผลการประเมินนี้เป็นการประเมินเบื้องต้น กรุณาตรวจสอบกับ RSO และเอกสาร IAEA SSR-6 ก่อนใช้งานจริง")

# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown("#### 🤖 ให้ AI ร่างแผนการขนส่ง")

    if not anthropic_key:
        st.warning("⚠️ กรุณาใส่ Anthropic API Key ใน sidebar เพื่อใช้งาน AI")
    else:
        if st.button("✨ ร่างแผนการขนส่ง", type="primary"):
            cat_str = f"ประเภท {class_result.get('category','—')}" if class_result.get("category") else "—"
            prompt = f"""คุณเป็นผู้เชี่ยวชาญด้านความปลอดภัยทางรังสีและกฎหมายพลังงานนิวเคลียร์ของไทย
จงร่างแผนการขนส่งวัสดุกัมมันตรังสีที่สมบูรณ์ตามข้อกำหนดของสำนักงานปรมาณูเพื่อสันติ (ปส.) และ IAEA SSR-6

ข้อมูล:
- องค์กร: {org_name} | ใบอนุญาต: {license_no}
- ที่อยู่: {org_address} | โทร: {org_phone}
- RSO: {rso_name} โทร. {rso_phone} | วันที่: {plan_date}
- นิวไคลด์: {isotope} | Activity: {activity} {unit} ({activity_tbq:.3e} TBq) ณ {act_date}
- ประเภทวัสดุ: {cat_str} | Special Form: {'ใช่' if special_form else 'ไม่ใช่'}
- Package Type: {pkg_result['package_type']} | UN Number: {pkg_result['un_number']}
- Proper Shipping Name: {pkg_result['proper_name']}
- Transport Index (TI): {ti} | Label: {label_info['label']} ({label_info['category']})
- รูปแบบขนส่ง: {transport_mode} | ต้นทาง: {origin} → ปลายทาง: {destination}
- ผู้ขนส่ง: {transporter} โทร. {trans_phone} | ยานพาหนะ: {vehicle} ทะเบียน {plate}
- ฉุกเฉิน: {emg1} / {emg2}

ร่างแผนภาษาไทยอย่างเป็นทางการ ประกอบด้วย:
1. วัตถุประสงค์และขอบเขต
2. ข้อมูลองค์กรและผู้รับผิดชอบ
3. ข้อมูลวัสดุกัมมันตรังสี (รวมการจำแนกประเภทและ Package Type)
4. มาตรการความปลอดภัยในการขนส่ง (ปรับตาม Package Type และ Category)
5. ข้อกำหนด Label และ Placard ที่ต้องติด
6. ข้อกำหนดยานพาหนะและผู้ขนส่ง
7. การตรวจสอบก่อนการขนส่ง
8. แผนฉุกเฉินระหว่างการขนส่ง
9. หน่วยงานที่แจ้งในกรณีฉุกเฉิน
10. เอกสารที่ต้องพาไปด้วย"""

            with st.spinner("AI กำลังร่างแผน..."):
                try:
                    resp = requests.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={"Content-Type": "application/json",
                                 "x-api-key": anthropic_key,
                                 "anthropic-version": "2023-06-01"},
                        json={"model": "claude-sonnet-4-20250514", "max_tokens": 2000,
                              "messages": [{"role": "user", "content": prompt}]},
                        timeout=60,
                    )
                    data = resp.json()
                    draft = data["content"][0]["text"]
                    st.session_state["transport_draft"] = draft
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {e}")

    if "transport_draft" in st.session_state:
        st.warning("⚠️ กรุณาตรวจสอบและแก้ไขก่อนใช้งานจริง")
        edited = st.text_area("แผนการขนส่ง (แก้ไขได้)",
                               value=st.session_state["transport_draft"], height=500)
        st.session_state["transport_draft_edited"] = edited

# ══════════════════════════════════════════════════════════════════════════════
with tab3:
    st.markdown("#### 📄 Export เป็น Word (.docx)")

    if st.button("📥 สร้างและดาวน์โหลด Word", type="primary"):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io

            doc = Document()

            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("แผนการขนส่งวัสดุกัมมันตรังสี")
            run.bold = True; run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1a, 0x4a, 0x7a)
            doc.add_paragraph("จัดทำตามข้อกำหนดสำนักงานปรมาณูเพื่อสันติ และ IAEA SSR-6 (Rev.1) 2018"
                               ).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            def add_table(doc, rows_data):
                t = doc.add_table(rows=0, cols=2)
                t.style = "Table Grid"
                for label, val in rows_data:
                    r = t.add_row().cells
                    r[0].text = label; r[1].text = str(val)
                return t

            # Section 1
            doc.add_heading("1. ข้อมูลองค์กรและผู้รับผิดชอบ", level=1)
            add_table(doc, [
                ("ชื่อองค์กร", org_name), ("เลขที่ใบอนุญาต", license_no),
                ("ที่อยู่", org_address), ("โทรศัพท์", org_phone),
                ("RSO", rso_name), ("โทรศัพท์ RSO", rso_phone),
                ("วันที่จัดทำแผน", str(plan_date)),
            ])
            doc.add_paragraph()

            # Section 2
            doc.add_heading("2. ข้อมูลวัสดุกัมมันตรังสี", level=1)
            cat_str = f"ประเภท {class_result.get('category','—')}"
            add_table(doc, [
                ("ชนิดนิวไคลด์", isotope),
                ("ค่ากัมมันตภาพ", f"{activity} {unit}  ({activity_tbq:.3e} TBq)"),
                ("วันที่อ้างอิง", str(act_date)),
                ("สถานะทางกายภาพ", phys_form),
                ("Special Form", "ใช่" if special_form else "ไม่ใช่"),
                ("ประเภทวัสดุกัมมันตรังสี", cat_str),
            ])
            doc.add_paragraph()

            # Section 3
            doc.add_heading("3. การจำแนกบรรจุภัณฑ์และ UN Number (IAEA SSR-6)", level=1)
            add_table(doc, [
                ("Package Type", pkg_result['package_type']),
                ("UN Number", pkg_result['un_number']),
                ("Proper Shipping Name", pkg_result['proper_name']),
                ("Transport Index (TI)", str(ti)),
                ("Label ที่ต้องติด", label_info['label']),
                ("Category", label_info['category']),
                ("เกณฑ์การจำแนก", pkg_result['detail']),
            ])
            doc.add_paragraph()

            # Section 4
            doc.add_heading("4. ข้อมูลการขนส่ง", level=1)
            add_table(doc, [
                ("รูปแบบการขนส่ง", transport_mode),
                ("ต้นทาง", origin), ("ปลายทาง", destination),
                ("ผู้ขนส่ง", transporter), ("โทรศัพท์ผู้ขนส่ง", trans_phone),
                ("ยานพาหนะ", vehicle), ("ทะเบียน", plate),
                ("ผู้ติดต่อฉุกเฉิน 1", emg1), ("ผู้ติดต่อฉุกเฉิน 2", emg2),
            ])
            doc.add_paragraph()

            # Section 5: AI draft
            draft_text = st.session_state.get("transport_draft_edited",
                          st.session_state.get("transport_draft", ""))
            if draft_text:
                doc.add_heading("5. รายละเอียดแผนการขนส่ง", level=1)
                for line in draft_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

            doc.add_paragraph()
            doc.add_paragraph("ลงนาม: ________________________________  วันที่: _______________")
            doc.add_paragraph("(ผู้รับใบอนุญาต / เจ้าหน้าที่ความปลอดภัยทางรังสี)")

            buf = io.BytesIO()
            doc.save(buf); buf.seek(0)

            st.download_button(
                label="⬇️ ดาวน์โหลด transport_plan.docx",
                data=buf, file_name="transport_plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.success("✅ ไฟล์ Word พร้อมดาวน์โหลดแล้วค่ะ")
        except ImportError:
            st.error("กรุณาติดตั้ง python-docx ก่อน: `pip install python-docx`")
        except Exception as e:
            st.error(f"เกิดข้อผิดพลาด: {e}")
