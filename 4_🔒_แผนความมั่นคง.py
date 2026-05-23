"""
🔒 Page 4 — แผนรักษาความมั่นคงปลอดภัยทางรังสี
อ้างอิง:
  - กฎกระทรวงความมั่นคงปลอดภัยทางรังสี พ.ศ. 2561
  - แนวทางการเขียนแผนรักษาความมั่นคงปลอดภัย (เอกสาร ปส.)
  - IAEA Nuclear Security Series
"""
import streamlit as st
import sys, os, requests
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from utils.data import D_VALUES_TBQ, ISOTOPE_INFO, UNIT_TO_TBQ, classify_material, USE_TYPE_CATEGORY

st.set_page_config(page_title="แผนรักษาความมั่นคงปลอดภัย", page_icon="🔒", layout="wide")
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@400;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Sarabun', sans-serif !important; }
.section-card {
    background: #f8fafc; border-radius: 10px; padding: 16px;
    border-left: 4px solid; margin: 8px 0;
}
</style>
""", unsafe_allow_html=True)

# ── ข้อมูลอ้างอิง ─────────────────────────────────────────────────────────────

# ระดับการรักษาความมั่นคงปลอดภัยตามประเภทวัสดุ
# อ้างอิง: กฎกระทรวงความมั่นคงปลอดภัยทางรังสี พ.ศ. 2561
SECURITY_LEVELS = {
    1: {
        "level": "ระดับ A (สูงสุด)",
        "color": "#dc2626",
        "bg": "#fef2f2",
        "desc": "วัสดุกัมมันตรังสีประเภท 1 — ต้องการมาตรการรักษาความมั่นคงสูงสุด",
        "measures": [
            "ระบบตรวจจับการบุกรุก (IDS) ที่เชื่อมต่อกับศูนย์เฝ้าระวัง 24 ชม.",
            "กล้อง CCTV ครอบคลุมทุกจุดเข้า-ออก",
            "การ์ดรักษาความปลอดภัยประจำการตลอด 24 ชั่วโมง",
            "ระบบควบคุมการเข้าถึง 2 ปัจจัย (Two-factor authentication)",
            "ประตูกั้นพื้นที่สองชั้น (Double barrier)",
            "แจ้ง ปส. ภายใน 1 ชั่วโมงเมื่อเกิดเหตุ",
        ]
    },
    2: {
        "level": "ระดับ B (สูง)",
        "color": "#ea580c",
        "bg": "#fff7ed",
        "desc": "วัสดุกัมมันตรังสีประเภท 2 — ต้องการมาตรการรักษาความมั่นคงสูง",
        "measures": [
            "ระบบตรวจจับการบุกรุก (IDS)",
            "กล้อง CCTV บริเวณจัดเก็บวัสดุฯ",
            "ระบบควบคุมการเข้าถึงด้วยบัตร/กุญแจ",
            "ตรวจสอบบัญชีวัสดุฯ อย่างน้อยสัปดาห์ละครั้ง",
            "แจ้ง ปส. ภายใน 4 ชั่วโมงเมื่อเกิดเหตุ",
        ]
    },
    3: {
        "level": "ระดับ C (กลาง)",
        "color": "#d97706",
        "bg": "#fffbeb",
        "desc": "วัสดุกัมมันตรังสีประเภท 3 — ต้องการมาตรการรักษาความมั่นคงระดับกลาง",
        "measures": [
            "ล็อกกุญแจสถานที่จัดเก็บ",
            "ควบคุมการเข้าถึงเฉพาะผู้ได้รับอนุญาต",
            "ตรวจสอบบัญชีวัสดุฯ อย่างน้อยเดือนละครั้ง",
            "แจ้ง ปส. ภายใน 24 ชั่วโมงเมื่อเกิดเหตุ",
        ]
    },
    4: {
        "level": "ระดับ D (พื้นฐาน)",
        "color": "#16a34a",
        "bg": "#f0fdf4",
        "desc": "วัสดุกัมมันตรังสีประเภท 4 และ 5 — มาตรการพื้นฐาน",
        "measures": [
            "จัดเก็บในที่ปลอดภัยมีกุญแจล็อก",
            "ตรวจสอบบัญชีวัสดุฯ อย่างน้อยปีละครั้ง",
            "แจ้ง ปส. ทันทีเมื่อเกิดเหตุ",
        ]
    },
}

# แมประเภทวัสดุ → ระดับความมั่นคง
def get_security_level(category: int) -> int:
    mapping = {1: 1, 2: 2, 3: 3, 4: 4, 5: 4}
    return mapping.get(category, 4)

# หัวข้อแผนรักษาความมั่นคงตามกฎกระทรวง
PLAN_SECTIONS = [
    ("1", "ที่มาและวัตถุประสงค์ของแผน",
     "อธิบายสาเหตุที่จัดทำ วัตถุประสงค์ ขอบเขต และกำหนดการทบทวนแผน"),
    ("2", "ขอบเขตของแผน",
     "ครอบคลุมทั้งทางกายภาพ (สถานที่ วัสดุกัมมันตรังสี) และทางข้อมูล"),
    ("3", "ข้อมูลรายละเอียดของสถานประกอบการ",
     "แผนที่ แผนผัง และรายละเอียดที่ตั้ง"),
    ("4", "ข้อมูลวัสดุกัมมันตรังสี",
     "ประเภท ชนิด ค่ากัมมันตภาพ และระดับการรักษาความมั่นคงที่ต้องจัดให้มี"),
    ("5", "ข้อมูลเวลาทำการและเส้นทางการเข้าออก",
     "เวลาเปิด-ปิด เส้นทางเข้าออก วิธีการใช้ประโยชน์"),
    ("6", "บทบาทและหน้าที่ของบุคลากร",
     "แผนผังพร้อมบทบาทและหน้าที่เกี่ยวกับการรักษาความมั่นคง"),
    ("7", "แผนการฝึกอบรมและทดสอบคุณสมบัติ",
     "หลักสูตร ความถี่ ผู้เข้าอบรม และการประเมินผล"),
    ("8", "แผนการเผชิญเหตุความมั่นคงปลอดภัย",
     "บทบาท การสื่อสาร แผนสำรอง การรายงาน และการดำเนินการหลังเหตุ"),
]

# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 🏢 ข้อมูลองค์กร")
    org_name    = st.text_input("ชื่อองค์กร / ผู้รับใบอนุญาต",
                                 placeholder="บริษัท ABC จำกัด")
    license_no  = st.text_input("เลขที่ใบอนุญาต", placeholder="4I0001/63F")
    org_address = st.text_area("ที่อยู่สถานประกอบการ", height=70)
    org_phone   = st.text_input("โทรศัพท์", placeholder="02-xxx-xxxx")
    rso_name    = st.text_input("ชื่อ RSO", placeholder="นายปรมาณู เพื่อสันติ")
    rso_phone   = st.text_input("โทรศัพท์ RSO", placeholder="08x-xxx-xxxx")
    plan_date   = st.date_input("วันที่จัดทำแผน")
    review_freq = st.selectbox("ความถี่การทบทวนแผน",
                                ["ปีละ 1 ครั้ง (ขั้นต่ำตามกฎหมาย)",
                                 "ปีละ 2 ครั้ง", "ทุก 6 เดือน"])
    st.markdown("---")
    anthropic_key = st.text_input("🔑 Anthropic API Key", type="password")

st.markdown("# 🔒 แผนรักษาความมั่นคงปลอดภัยทางรังสี")
st.caption(
    "อ้างอิง: กฎกระทรวงความมั่นคงปลอดภัยทางรังสี พ.ศ. 2561 | "
    "แนวทางการเขียนแผนรักษาความมั่นคงปลอดภัย (ปส.) | IAEA Nuclear Security Series"
)

tab1, tab2, tab3, tab4 = st.tabs([
    "☢️ วัสดุและระดับความมั่นคง",
    "📋 กรอกข้อมูลแผน (8 หัวข้อ)",
    "🔍 ตรวจสอบความครบถ้วน",
    "🤖 AI ร่างแผน & Export",
])

isotope_list = list(D_VALUES_TBQ.keys())
use_list = ["— ไม่พบในตาราง / คำนวณ A/D —"] + list(USE_TYPE_CATEGORY.keys())

# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    st.subheader("ข้อมูลวัสดุกัมมันตรังสีและการกำหนดระดับความมั่นคง")
    st.caption(
        "ระดับความมั่นคงขึ้นอยู่กับประเภทวัสดุกัมมันตรังสี "
        "ตามกฎกระทรวงความมั่นคงปลอดภัยทางรังสี พ.ศ. 2561"
    )

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### ☢️ ข้อมูลวัสดุกัมมันตรังสี")
        isotope  = st.selectbox("ชนิดนิวไคลด์", isotope_list,
                                 index=isotope_list.index("Ir-192"))
        ca, cb   = st.columns(2)
        activity = ca.number_input("ค่ากัมมันตภาพ", min_value=0.0,
                                    value=5.5, format="%.4f")
        unit     = cb.selectbox("หน่วย", list(UNIT_TO_TBQ.keys()),
                                 index=list(UNIT_TO_TBQ.keys()).index("Ci"))
        use_type = st.selectbox("การใช้ประโยชน์", use_list)
        activity_tbq = activity * UNIT_TO_TBQ.get(unit, 1)
        st.caption(f"= **{activity_tbq:.4e} TBq**")

        if isotope in ISOTOPE_INFO:
            info = ISOTOPE_INFO[isotope]
            st.info(f"**{isotope}** | ครึ่งชีวิต: {info['halfLife']} | "
                    f"UN: {info['unNo']}")

        # จำแนกประเภทวัสดุ
        ut = use_type if use_type != "— ไม่พบในตาราง / คำนวณ A/D —" else ""
        class_result = classify_material(isotope, activity_tbq, ut)

        from utils.data import CATEGORY_COLORS, CATEGORY_BG
        if class_result.get("category"):
            cat = class_result["category"]
            st.markdown(f"""
            <div style="background:{CATEGORY_BG[cat]};border:2px solid {CATEGORY_COLORS[cat]};
                        border-radius:10px;padding:12px 16px;margin-top:8px;">
                <b style="color:{CATEGORY_COLORS[cat]};font-size:16px;">ประเภทวัสดุ: ประเภท {cat}</b><br>
                <small style="color:#6b7280;">{class_result.get('detail','')}</small>
            </div>
            """, unsafe_allow_html=True)

        # เพิ่มแหล่งรังสีหลายชนิด
        st.markdown("#### ➕ วัสดุกัมมันตรังสีที่ครอบครองทั้งหมด")
        if "sources_security" not in st.session_state:
            st.session_state.sources_security = []

        if st.button("เพิ่มวัสดุกัมมันตรังสี"):
            st.session_state.sources_security.append({
                "isotope": isotope, "activity": activity,
                "unit": unit, "use_type": use_type,
                "category": class_result.get("category"),
                "activity_tbq": activity_tbq,
            })
            st.rerun()

        if st.session_state.sources_security:
            import pandas as pd
            rows = []
            for i, s in enumerate(st.session_state.sources_security):
                rows.append({
                    "#": i+1,
                    "Isotope": s["isotope"],
                    "Activity": f"{s['activity']} {s['unit']}",
                    "ประเภท": f"ประเภท {s['category']}" if s['category'] else "—",
                })
            st.dataframe(pd.DataFrame(rows), hide_index=True,
                         use_container_width=True)
            if st.button("ล้างรายการ"):
                st.session_state.sources_security = []
                st.rerun()

    with c2:
        st.markdown("#### 🔒 ระดับความมั่นคงปลอดภัยที่ต้องจัดให้มี")

        # หาระดับสูงสุดจากวัสดุทั้งหมด
        all_cats = [s["category"] for s in st.session_state.sources_security
                    if s.get("category")]
        if class_result.get("category"):
            all_cats.append(class_result["category"])

        sec_level = get_security_level(min(all_cats)) if all_cats else 4
        sec_info  = SECURITY_LEVELS[sec_level]

        st.markdown(f"""
        <div style="background:{sec_info['bg']};border:3px solid {sec_info['color']};
                    border-radius:14px;padding:20px;margin-bottom:16px;">
            <div style="display:flex;align-items:center;gap:12px;margin-bottom:12px;">
                <span style="font-size:28px;">🔒</span>
                <div>
                    <b style="font-size:18px;color:{sec_info['color']};">
                        {sec_info['level']}</b><br>
                    <small style="color:#6b7280;">{sec_info['desc']}</small>
                </div>
            </div>
            <hr style="border-color:{sec_info['color']}44;margin:10px 0;">
            <p style="margin:0 0 8px;font-size:13px;font-weight:600;color:#374151;">
                มาตรการที่ต้องจัดให้มี:</p>
        </div>
        """, unsafe_allow_html=True)

        for m in sec_info["measures"]:
            st.markdown(f"✅ {m}")

        st.divider()
        # แสดงระดับทั้งหมดเพื่ออ้างอิง
        with st.expander("📋 ดูระดับความมั่นคงทั้งหมด"):
            for lv, info in SECURITY_LEVELS.items():
                st.markdown(f"""
                <div style="background:{info['bg']};border-left:4px solid {info['color']};
                            border-radius:0 8px 8px 0;padding:10px 14px;margin:6px 0;">
                    <b style="color:{info['color']};">{info['level']}</b>
                    <span style="font-size:12px;color:#6b7280;"> — {info['desc']}</span>
                </div>
                """, unsafe_allow_html=True)

        # บันทึก session
        st.session_state["security"] = {
            "isotope": isotope, "activity": activity, "unit": unit,
            "activity_tbq": activity_tbq, "use_type": use_type,
            "class_result": class_result,
            "sec_level": sec_level, "sec_info": sec_info,
            "all_sources": st.session_state.sources_security,
        }

# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.subheader("กรอกข้อมูลแผน 8 หัวข้อ ตามกฎกระทรวง พ.ศ. 2561")
    st.caption("ทุกหัวข้อต้องครบถ้วนเพื่อประกอบการยื่นขอใบอนุญาตต่อ ปส.")

    if "plan_data" not in st.session_state:
        st.session_state["plan_data"] = {}

    plan = st.session_state["plan_data"]

    # ── หัวข้อ 1: วัตถุประสงค์ ──────────────────────────────────────────────
    with st.expander("📌 หัวข้อที่ 1 — ที่มาและวัตถุประสงค์ของแผน", expanded=True):
        plan["objective"] = st.text_area(
            "วัตถุประสงค์และขอบเขตแผน",
            value=plan.get("objective",
                f"แผนรักษาความมั่นคงปลอดภัยฉบับนี้จัดทำขึ้นโดย{org_name or '[ชื่อองค์กร]'} "
                "เพื่อให้สามารถปกป้องวัสดุกัมมันตรังสีได้อย่างเพียงพอจากการเข้าถึงโดยไม่ได้รับ"
                "อนุญาต การโจรกรรม และการก่อวินาศกรรม ให้สอดคล้องกับกฎกระทรวงความมั่นคง"
                "ปลอดภัยทางรังสี พ.ศ. 2561 และมีการทบทวนทุก 1 ปี"),
            height=100, key="p1"
        )

    # ── หัวข้อ 2: ขอบเขต ────────────────────────────────────────────────────
    with st.expander("📌 หัวข้อที่ 2 — ขอบเขตของแผน"):
        plan["scope_physical"] = st.text_area(
            "ขอบเขตทางกายภาพ (สถานที่ อาคาร พื้นที่)",
            value=plan.get("scope_physical", ""),
            height=80, key="p2a"
        )
        plan["scope_info"] = st.text_area(
            "ขอบเขตด้านข้อมูล (เอกสาร ระบบสารสนเทศ ข้อมูลวัสดุฯ)",
            value=plan.get("scope_info", ""),
            height=80, key="p2b"
        )

    # ── หัวข้อ 3: รายละเอียดสถานประกอบการ ───────────────────────────────────
    with st.expander("📌 หัวข้อที่ 3 — ข้อมูลรายละเอียดสถานประกอบการ"):
        c3a, c3b = st.columns(2)
        plan["facility_address"] = c3a.text_input(
            "ที่อยู่สถานประกอบการ",
            value=plan.get("facility_address", org_address), key="p3a")
        plan["facility_phone"] = c3b.text_input(
            "โทรศัพท์สถานประกอบการ",
            value=plan.get("facility_phone", org_phone), key="p3b")
        plan["nearby_facilities"] = st.text_area(
            "สถานที่สำคัญโดยรอบ (โรงพยาบาล สถานีตำรวจ สถานีดับเพลิง)",
            value=plan.get("nearby_facilities", ""), height=70, key="p3c"
        )
        plan["police_distance"] = st.text_input(
            "สถานีตำรวจที่ใกล้ที่สุด (ชื่อ + ระยะทาง)",
            value=plan.get("police_distance", ""), key="p3d"
        )

    # ── หัวข้อ 4: ข้อมูลวัสดุกัมมันตรังสี ───────────────────────────────────
    with st.expander("📌 หัวข้อที่ 4 — ข้อมูลวัสดุกัมมันตรังสีและระดับความมั่นคง"):
        sec = st.session_state.get("security", {})
        cr  = sec.get("class_result", {})
        st.info(
            f"ระบบดึงข้อมูลจาก Tab 1 อัตโนมัติ: "
            f"**{sec.get('isotope','—')}** ประเภท {cr.get('category','—')} → "
            f"ระดับความมั่นคง **{SECURITY_LEVELS.get(sec.get('sec_level',4),{}).get('level','—')}**"
        )
        plan["source_storage"] = st.text_area(
            "สถานที่จัดเก็บวัสดุกัมมันตรังสี (อาคาร ห้อง ตู้)",
            value=plan.get("source_storage", ""), height=70, key="p4a"
        )
        plan["access_control_device"] = st.text_input(
            "อุปกรณ์ควบคุมการเข้าถึง (กุญแจ, บัตร, รหัส, IDS, CCTV)",
            value=plan.get("access_control_device", ""), key="p4b"
        )
        plan["inventory_freq"] = st.selectbox(
            "ความถี่การตรวจสอบบัญชีวัสดุฯ",
            ["ทุกวัน", "ทุกสัปดาห์", "ทุกเดือน", "ทุกไตรมาส", "ทุกปี"],
            key="p4c"
        )

    # ── หัวข้อ 5: เวลาทำการและเส้นทาง ───────────────────────────────────────
    with st.expander("📌 หัวข้อที่ 5 — ข้อมูลเวลาทำการและเส้นทางการเข้าออก"):
        c5a, c5b = st.columns(2)
        plan["working_hours"] = c5a.text_input(
            "เวลาทำการปกติ", value=plan.get("working_hours", "08:00–17:00 น."),
            key="p5a"
        )
        plan["access_routes"] = c5b.text_input(
            "เส้นทางเข้า-ออกหลัก", value=plan.get("access_routes", ""),
            key="p5b"
        )
        plan["after_hours"] = st.text_area(
            "มาตรการนอกเวลาทำการ (กลางคืน วันหยุด)",
            value=plan.get("after_hours",
                "ล็อกกุญแจสองชั้น เปิดระบบ IDS แจ้งเจ้าหน้าที่รักษาความปลอดภัย"),
            height=70, key="p5c"
        )

    # ── หัวข้อ 6: บุคลากร ────────────────────────────────────────────────────
    with st.expander("📌 หัวข้อที่ 6 — บทบาทและหน้าที่ของบุคลากร"):
        st.markdown("**ผู้มีหน้าที่รับผิดชอบด้านความมั่นคงปลอดภัย**")

        if "personnel" not in st.session_state:
            st.session_state.personnel = [
                {"name": org_name or "ผู้รับใบอนุญาต", "position": "ผู้รับใบอนุญาต",
                 "role": "กำหนดนโยบายและรับผิดชอบสูงสุด", "phone": org_phone or ""},
                {"name": rso_name or "RSO", "position": "เจ้าหน้าที่ความปลอดภัยทางรังสี",
                 "role": "ควบคุมดูแลการใช้งานและความมั่นคงปลอดภัย", "phone": rso_phone or ""},
            ]

        personnel_updated = []
        for i, p in enumerate(st.session_state.personnel):
            pc1, pc2, pc3, pc4, pc5 = st.columns([2, 2, 3, 2, 0.5])
            pname = pc1.text_input(f"ชื่อ #{i+1}", value=p["name"], key=f"pn_{i}")
            ppos  = pc2.text_input(f"ตำแหน่ง #{i+1}", value=p["position"], key=f"pp_{i}")
            prole = pc3.text_input(f"หน้าที่ #{i+1}", value=p["role"], key=f"pr_{i}")
            pph   = pc4.text_input(f"โทรศัพท์ #{i+1}", value=p["phone"], key=f"pt_{i}")
            pc5.markdown("<br>", unsafe_allow_html=True)
            if not pc5.button("🗑️", key=f"pdel_{i}"):
                personnel_updated.append({
                    "name": pname, "position": ppos,
                    "role": prole, "phone": pph
                })
        st.session_state.personnel = personnel_updated

        if st.button("➕ เพิ่มบุคลากร"):
            st.session_state.personnel.append({
                "name": "", "position": "", "role": "", "phone": ""})
            st.rerun()

        plan["personnel"] = st.session_state.personnel

    # ── หัวข้อ 7: การฝึกอบรม ─────────────────────────────────────────────────
    with st.expander("📌 หัวข้อที่ 7 — แผนการฝึกอบรมและทดสอบคุณสมบัติ"):
        st.markdown("**หลักสูตรการฝึกอบรมด้านความมั่นคงปลอดภัย**")

        if "training" not in st.session_state:
            st.session_state.training = [
                {"course": "ความปลอดภัยทางรังสีสำหรับผู้ปฏิบัติงานใหม่",
                 "freq": "ก่อนเริ่มงาน", "target": "ผู้ปฏิบัติงานใหม่", "format": "ภายใน"},
                {"course": "การรักษาความมั่นคงปลอดภัยทางรังสี",
                 "freq": "ปีละ 1 ครั้ง", "target": "ผู้ปฏิบัติงานทุกคน", "format": "ภายใน/ภายนอก"},
                {"course": "การเผชิญเหตุฉุกเฉินทางรังสี",
                 "freq": "ปีละ 1 ครั้ง", "target": "ผู้ปฏิบัติงานและ RSO", "format": "ซ้อมแผน"},
            ]

        training_updated = []
        for i, t in enumerate(st.session_state.training):
            tc1, tc2, tc3, tc4, tc5 = st.columns([3, 2, 2, 1.5, 0.5])
            tcourse = tc1.text_input(f"หลักสูตร #{i+1}", value=t["course"], key=f"tc_{i}")
            tfreq   = tc2.text_input(f"ความถี่ #{i+1}", value=t["freq"], key=f"tf_{i}")
            ttarg   = tc3.text_input(f"ผู้เข้าอบรม #{i+1}", value=t["target"], key=f"tt_{i}")
            tfmt    = tc4.selectbox(f"รูปแบบ #{i+1}",
                                     ["ภายใน", "ภายนอก", "ออนไลน์", "ซ้อมแผน"],
                                     index=["ภายใน","ภายนอก","ออนไลน์","ซ้อมแผน"].index(
                                         t["format"]) if t["format"] in
                                         ["ภายใน","ภายนอก","ออนไลน์","ซ้อมแผน"] else 0,
                                     key=f"fm_{i}")
            tc5.markdown("<br>", unsafe_allow_html=True)
            if not tc5.button("🗑️", key=f"tdel_{i}"):
                training_updated.append({
                    "course": tcourse, "freq": tfreq,
                    "target": ttarg, "format": tfmt
                })
        st.session_state.training = training_updated

        if st.button("➕ เพิ่มหลักสูตร"):
            st.session_state.training.append({
                "course": "", "freq": "", "target": "", "format": "ภายใน"})
            st.rerun()

        plan["training"] = st.session_state.training

    # ── หัวข้อ 8: แผนเผชิญเหตุ ───────────────────────────────────────────────
    with st.expander("📌 หัวข้อที่ 8 — แผนการเผชิญเหตุความมั่นคงปลอดภัย"):
        plan["incident_roles"] = st.text_area(
            "ก) บทบาทและหน้าที่ในการเผชิญเหตุ",
            value=plan.get("incident_roles",
                "RSO: ประเมินสถานการณ์และสั่งการ | ผู้ปฏิบัติงาน: อพยพออกจากพื้นที่ | "
                "ผู้รับใบอนุญาต: แจ้งหน่วยงานภายนอก"),
            height=80, key="p8a"
        )
        plan["comm_method"] = st.text_area(
            "ข) วิธีการติดต่อสื่อสารในการเผชิญเหตุ",
            value=plan.get("comm_method",
                "ติดต่อ RSO: โทรศัพท์มือถือ | ติดต่อ ปส.: 02-596-7600 | "
                "นอกเวลา: 089-200-6243 | วิทยุสื่อสาร (สำรอง)"),
            height=80, key="p8b"
        )
        plan["contingency"] = st.text_area(
            "ค) แผนสำรองกรณีระบบขัดข้อง",
            value=plan.get("contingency", ""), height=70, key="p8c"
        )
        plan["threat_escalation"] = st.text_area(
            "ง) วิธีดำเนินการเมื่อภัยคุกคามยกระดับสูงขึ้น",
            value=plan.get("threat_escalation",
                "เพิ่มความถี่การตรวจสอบ | แจ้ง ปส. ทันที | "
                "ประสาน ตำรวจ / กรมป้องกันฯ ตามความจำเป็น"),
            height=80, key="p8d"
        )
        plan["post_incident"] = st.text_area(
            "จ) แผนการดำเนินการภายหลังเกิดเหตุ",
            value=plan.get("post_incident",
                "สำรวจความเสียหาย | บันทึกรายงานเหตุการณ์ | "
                "ส่งรายงานต่อ ปส. ภายใน 24 ชม. | ทบทวนและปรับปรุงแผน"),
            height=80, key="p8e"
        )

        # ผู้ติดต่อฉุกเฉิน
        st.markdown("**หน่วยงานที่ต้องแจ้งในกรณีเกิดเหตุ**")
        import pandas as pd
        emergency_contacts = pd.DataFrame([
            {"หน่วยงาน": "สำนักงานปรมาณูเพื่อสันติ (ปส.)",
             "โทรศัพท์": "02-596-7600", "หมายเหตุ": "ตลอด 24 ชม."},
            {"หน่วยงาน": "ปส. สายด่วนนอกเวลา",
             "โทรศัพท์": "089-200-6243", "หมายเหตุ": "ตลอด 24 ชม."},
            {"หน่วยงาน": "กรมป้องกันและบรรเทาสาธารณภัย",
             "โทรศัพท์": "1784", "หมายเหตุ": "กรณีสาธารณภัย"},
            {"หน่วยงาน": "ศูนย์นเรนทร กระทรวงสาธารณสุข",
             "โทรศัพท์": "1669", "หมายเหตุ": "กรณีมีผู้บาดเจ็บ"},
            {"หน่วยงาน": "สำนักงานตำรวจแห่งชาติ",
             "โทรศัพท์": "191", "หมายเหตุ": "กรณีก่อการร้าย"},
        ])
        st.dataframe(emergency_contacts, hide_index=True, use_container_width=True)

    st.session_state["plan_data"] = plan

# ══════════════════════════════════════════════════════════════════════════════
with tab3:
    st.subheader("🔍 ตรวจสอบความครบถ้วนของแผน")
    st.caption("ตรวจสอบว่าข้อมูลครบถ้วนตามข้อกำหนดของกฎกระทรวง")

    plan3 = st.session_state.get("plan_data", {})
    sec3  = st.session_state.get("security", {})

    # Checklist
    checks = [
        ("ข้อมูลองค์กร", bool(org_name and license_no and rso_name)),
        ("วัตถุประสงค์และขอบเขต (หัวข้อ 1-2)",
         bool(plan3.get("objective") and plan3.get("scope_physical"))),
        ("รายละเอียดสถานประกอบการ (หัวข้อ 3)",
         bool(plan3.get("facility_address"))),
        ("ข้อมูลวัสดุกัมมันตรังสีและระดับความมั่นคง (หัวข้อ 4)",
         bool(sec3.get("class_result", {}).get("category") and
              plan3.get("source_storage"))),
        ("เวลาทำการและเส้นทาง (หัวข้อ 5)",
         bool(plan3.get("working_hours"))),
        ("บุคลากรและหน้าที่ (หัวข้อ 6)",
         bool(st.session_state.get("personnel"))),
        ("แผนการฝึกอบรม (หัวข้อ 7)",
         bool(st.session_state.get("training"))),
        ("แผนเผชิญเหตุ (หัวข้อ 8)",
         bool(plan3.get("incident_roles") and plan3.get("comm_method"))),
    ]

    done  = sum(1 for _, v in checks if v)
    total = len(checks)
    pct   = done / total * 100

    # Progress
    color_pct = "#16a34a" if pct == 100 else ("#d97706" if pct >= 60 else "#dc2626")
    st.markdown(f"""
    <div style="background:#f8fafc;border-radius:12px;padding:18px;margin-bottom:16px;">
        <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;">
            <b style="font-size:15px;">ความครบถ้วนของแผน</b>
            <b style="font-size:20px;color:{color_pct};">{done}/{total} หัวข้อ ({pct:.0f}%)</b>
        </div>
        <div style="background:#e5e7eb;border-radius:20px;height:10px;">
            <div style="background:{color_pct};width:{pct}%;height:10px;
                        border-radius:20px;transition:width 0.3s;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Checklist items
    for label, ok in checks:
        icon = "✅" if ok else "❌"
        color = "#16a34a" if ok else "#dc2626"
        st.markdown(f'<span style="color:{color};font-size:14px;">{icon} {label}</span>',
                    unsafe_allow_html=True)

    if pct == 100:
        st.success("🎉 แผนครบถ้วนทุกหัวข้อ พร้อมให้ AI ร่างและ Export Word ค่ะ!")
    else:
        missing = [l for l, v in checks if not v]
        st.warning(f"⚠️ ยังขาดข้อมูล: {', '.join(missing)}")

    # แสดงสรุปข้อมูลทั้งหมด
    st.divider()
    st.markdown("#### 📋 สรุปข้อมูลสำหรับตรวจสอบ")
    if sec3.get("class_result", {}).get("category"):
        sl = sec3.get("sec_level", 4)
        si = SECURITY_LEVELS[sl]
        c_l, c_r = st.columns(2)
        c_l.metric("ประเภทวัสดุกัมมันตรังสี",
                   f"ประเภท {sec3['class_result']['category']}")
        c_r.metric("ระดับความมั่นคงปลอดภัย", si["level"])

# ══════════════════════════════════════════════════════════════════════════════
with tab4:
    st.subheader("🤖 AI ร่างแผนรักษาความมั่นคงปลอดภัย")

    plan4 = st.session_state.get("plan_data", {})
    sec4  = st.session_state.get("security", {})
    pers  = st.session_state.get("personnel", [])
    train = st.session_state.get("training", [])

    if not anthropic_key:
        st.warning("⚠️ กรุณาใส่ Anthropic API Key ใน sidebar")
    else:
        if st.button("✨ ร่างแผนรักษาความมั่นคงปลอดภัย", type="primary"):

            cr4 = sec4.get("class_result", {})
            sl4 = sec4.get("sec_level", 4)
            si4 = SECURITY_LEVELS.get(sl4, {})

            personnel_text = "\n".join(
                [f"  - {p['name']} ({p['position']}): {p['role']} โทร. {p['phone']}"
                 for p in pers]) if pers else "ยังไม่ระบุ"

            training_text = "\n".join(
                [f"  - {t['course']} | {t['freq']} | {t['target']} | {t['format']}"
                 for t in train]) if train else "ยังไม่ระบุ"

            sources_text = "\n".join(
                [f"  - {s['isotope']} {s['activity']} {s['unit']} → ประเภท {s['category']}"
                 for s in sec4.get("all_sources", [])]) or \
                f"  - {sec4.get('isotope','—')} {sec4.get('activity','—')} {sec4.get('unit','—')} → ประเภท {cr4.get('category','—')}"

            prompt = f"""คุณเป็นผู้เชี่ยวชาญด้านความปลอดภัยทางรังสีและกฎหมายพลังงานนิวเคลียร์ของไทย
จงร่างแผนรักษาความมั่นคงปลอดภัยทางรังสีที่สมบูรณ์ตามกฎกระทรวงความมั่นคงปลอดภัยทางรังสี พ.ศ. 2561

ข้อมูลองค์กร:
- องค์กร: {org_name} | ใบอนุญาต: {license_no}
- ที่อยู่: {org_address}
- RSO: {rso_name} โทร. {rso_phone} | วันที่: {plan_date}
- ความถี่ทบทวนแผน: {review_freq}

วัสดุกัมมันตรังสีที่ครอบครอง:
{sources_text}
- ระดับความมั่นคงปลอดภัยที่ต้องจัดให้มี: {si4.get('level','—')}

ข้อมูลที่กรอกไว้:
- วัตถุประสงค์: {plan4.get('objective','—')}
- ขอบเขตกายภาพ: {plan4.get('scope_physical','—')}
- ขอบเขตข้อมูล: {plan4.get('scope_info','—')}
- ที่อยู่สถานประกอบการ: {plan4.get('facility_address','—')}
- สถานที่สำคัญโดยรอบ: {plan4.get('nearby_facilities','—')}
- สถานีตำรวจใกล้สุด: {plan4.get('police_distance','—')}
- สถานที่จัดเก็บวัสดุฯ: {plan4.get('source_storage','—')}
- อุปกรณ์ควบคุมการเข้าถึง: {plan4.get('access_control_device','—')}
- ความถี่ตรวจสอบบัญชีวัสดุฯ: {plan4.get('inventory_freq','—')}
- เวลาทำการ: {plan4.get('working_hours','—')}
- มาตรการนอกเวลา: {plan4.get('after_hours','—')}

บุคลากร:
{personnel_text}

แผนฝึกอบรม:
{training_text}

แผนเผชิญเหตุ:
- บทบาท: {plan4.get('incident_roles','—')}
- การสื่อสาร: {plan4.get('comm_method','—')}
- แผนสำรอง: {plan4.get('contingency','—')}
- การยกระดับภัยคุกคาม: {plan4.get('threat_escalation','—')}
- หลังเกิดเหตุ: {plan4.get('post_incident','—')}

มาตรการระดับ {si4.get('level','—')} ที่ต้องมี:
{chr(10).join(['- ' + m for m in si4.get('measures',[])])}

จงร่างแผนรักษาความมั่นคงปลอดภัยทางรังสีภาษาไทยอย่างเป็นทางการ
ครอบคลุม 8 องค์ประกอบหลักตามกฎกระทรวง:
1. บทนำ — วัตถุประสงค์ ขอบเขต และกำหนดการทบทวน
2. รายละเอียดสถานประกอบการ
3. ข้อมูลการบริหารจัดการความมั่นคงปลอดภัย (หน้าที่ การฝึกอบรม การเข้าถึง การปกป้องข้อมูล)
4. รายละเอียดระบบรักษาความมั่นคงปลอดภัย (ภัยคุกคาม การประเมิน รูปแบบ มาตรการ)
5. แนวปฏิบัติการรักษาความมั่นคงปลอดภัย (เวลาทำการ กุญแจ บัญชีวัสดุฯ)
6. การเผชิญเหตุความมั่นคงปลอดภัย
7. อ้างอิงและภาคผนวก"""

            with st.spinner("AI กำลังร่างแผน..."):
                try:
                    resp = requests.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={"Content-Type": "application/json",
                                 "x-api-key": anthropic_key,
                                 "anthropic-version": "2023-06-01"},
                        json={"model": "claude-sonnet-4-20250514",
                              "max_tokens": 2000,
                              "messages": [{"role": "user", "content": prompt}]},
                        timeout=60,
                    )
                    data = resp.json()
                    draft = data["content"][0]["text"]
                    st.session_state["security_draft"] = draft
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {e}")

    if "security_draft" in st.session_state:
        st.warning("⚠️ กรุณาตรวจสอบและแก้ไขก่อนใช้งานจริง")
        edited = st.text_area("แผนรักษาความมั่นคงปลอดภัย (แก้ไขได้โดยตรง)",
                               value=st.session_state["security_draft"], height=500)
        st.session_state["security_draft_edited"] = edited

        st.divider()
        if st.button("📥 Export Word (.docx)", type="primary"):
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io, pandas as pd

                doc = Document()

                # Title
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = title.add_run("แผนรักษาความมั่นคงปลอดภัยทางรังสี")
                r.bold = True; r.font.size = Pt(18)
                r.font.color.rgb = RGBColor(0x1a, 0x4a, 0x7a)
                sub = doc.add_paragraph()
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub.add_run(
                    "จัดทำตามกฎกระทรวงความมั่นคงปลอดภัยทางรังสี พ.ศ. 2561"
                ).italic = True
                doc.add_paragraph()

                def add_table(doc, rows_data):
                    t = doc.add_table(rows=0, cols=2)
                    t.style = "Table Grid"
                    for lbl, val in rows_data:
                        row = t.add_row().cells
                        row[0].text = lbl; row[1].text = str(val)

                # ส่วนที่ 1: ข้อมูลองค์กร
                doc.add_heading("1. ข้อมูลองค์กรและผู้รับผิดชอบ", level=1)
                add_table(doc, [
                    ("ชื่อองค์กร", org_name),
                    ("เลขที่ใบอนุญาต", license_no),
                    ("ที่อยู่", org_address),
                    ("โทรศัพท์", org_phone),
                    ("RSO", rso_name),
                    ("โทรศัพท์ RSO", rso_phone),
                    ("วันที่จัดทำ", str(plan_date)),
                    ("ความถี่การทบทวนแผน", review_freq),
                ])
                doc.add_paragraph()

                # ส่วนที่ 2: วัสดุและระดับความมั่นคง
                doc.add_heading("2. ข้อมูลวัสดุกัมมันตรังสีและระดับความมั่นคง", level=1)
                cr4 = sec4.get("class_result", {})
                sl4 = sec4.get("sec_level", 4)
                si4 = SECURITY_LEVELS.get(sl4, {})
                add_table(doc, [
                    ("ชนิดนิวไคลด์", sec4.get("isotope", "—")),
                    ("ค่ากัมมันตภาพ",
                     f"{sec4.get('activity','—')} {sec4.get('unit','—')} "
                     f"({sec4.get('activity_tbq',0):.3e} TBq)"),
                    ("ประเภทวัสดุกัมมันตรังสี",
                     f"ประเภท {cr4.get('category','—')}"),
                    ("ระดับความมั่นคงปลอดภัย", si4.get("level", "—")),
                    ("สถานที่จัดเก็บ", plan4.get("source_storage", "—")),
                    ("อุปกรณ์ควบคุมการเข้าถึง",
                     plan4.get("access_control_device", "—")),
                    ("ความถี่ตรวจสอบบัญชีวัสดุฯ",
                     plan4.get("inventory_freq", "—")),
                ])
                doc.add_paragraph()

                # ส่วนที่ 3: บุคลากร
                if pers:
                    doc.add_heading("3. บุคลากรและหน้าที่รับผิดชอบ", level=1)
                    t2 = doc.add_table(rows=0, cols=4)
                    t2.style = "Table Grid"
                    hdr = t2.add_row().cells
                    hdr[0].text = "ชื่อ"; hdr[1].text = "ตำแหน่ง"
                    hdr[2].text = "หน้าที่"; hdr[3].text = "โทรศัพท์"
                    for p in pers:
                        row = t2.add_row().cells
                        row[0].text = p["name"]; row[1].text = p["position"]
                        row[2].text = p["role"]; row[3].text = p["phone"]
                    doc.add_paragraph()

                # ส่วนที่ 4: การฝึกอบรม
                if train:
                    doc.add_heading("4. แผนการฝึกอบรม", level=1)
                    t3 = doc.add_table(rows=0, cols=4)
                    t3.style = "Table Grid"
                    hdr3 = t3.add_row().cells
                    hdr3[0].text = "หลักสูตร"; hdr3[1].text = "ความถี่"
                    hdr3[2].text = "ผู้เข้าอบรม"; hdr3[3].text = "รูปแบบ"
                    for t_ in train:
                        row3 = t3.add_row().cells
                        row3[0].text = t_["course"]; row3[1].text = t_["freq"]
                        row3[2].text = t_["target"]; row3[3].text = t_["format"]
                    doc.add_paragraph()

                # ส่วนที่ 5: แผนเผชิญเหตุ
                doc.add_heading("5. แผนการเผชิญเหตุความมั่นคงปลอดภัย", level=1)
                add_table(doc, [
                    ("บทบาทและหน้าที่", plan4.get("incident_roles", "—")),
                    ("การสื่อสาร", plan4.get("comm_method", "—")),
                    ("แผนสำรอง", plan4.get("contingency", "—")),
                    ("การยกระดับภัยคุกคาม", plan4.get("threat_escalation", "—")),
                    ("การดำเนินการหลังเกิดเหตุ", plan4.get("post_incident", "—")),
                ])
                doc.add_paragraph()

                # ส่วนที่ 6: หน่วยงานฉุกเฉิน
                doc.add_heading("6. หน่วยงานที่ต้องแจ้งในกรณีเกิดเหตุ", level=1)
                t4 = doc.add_table(rows=0, cols=3)
                t4.style = "Table Grid"
                hdr4 = t4.add_row().cells
                hdr4[0].text = "หน่วยงาน"; hdr4[1].text = "โทรศัพท์"
                hdr4[2].text = "หมายเหตุ"
                for row_data in [
                    ("ปส.", "02-596-7600 / 089-200-6243", "ตลอด 24 ชม."),
                    ("กรมป้องกันฯ", "1784", "กรณีสาธารณภัย"),
                    ("ศูนย์นเรนทร", "1669", "กรณีมีผู้บาดเจ็บ"),
                    ("ตำรวจ", "191", "กรณีก่อการร้าย"),
                ]:
                    row4 = t4.add_row().cells
                    row4[0].text = row_data[0]; row4[1].text = row_data[1]
                    row4[2].text = row_data[2]
                doc.add_paragraph()

                # ส่วนที่ 7: AI Draft
                doc.add_heading("7. รายละเอียดแผนรักษาความมั่นคงปลอดภัย (AI Draft)",
                                level=1)
                draft_f = st.session_state.get(
                    "security_draft_edited",
                    st.session_state.get("security_draft", ""))
                for line in draft_f.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

                doc.add_paragraph()
                doc.add_paragraph(
                    "ลงนาม: ________________________________  วันที่: _______________")
                doc.add_paragraph(
                    "(ผู้รับใบอนุญาต / เจ้าหน้าที่ความปลอดภัยทางรังสี)")

                buf = io.BytesIO(); doc.save(buf); buf.seek(0)
                st.download_button(
                    label="⬇️ ดาวน์โหลด security_plan.docx",
                    data=buf, file_name="security_plan.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("✅ ไฟล์ Word พร้อมดาวน์โหลดแล้วค่ะ")

            except ImportError:
                st.error("กรุณาติดตั้ง python-docx: `pip install python-docx`")
            except Exception as e:
                st.error(f"เกิดข้อผิดพลาด: {e}")
